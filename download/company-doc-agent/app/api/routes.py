"""
FastAPI 路由定义
提供 REST API 接口供前端和外部调用

优化：
1. 新增 SSE 流式对话接口 /chat/stream — 实时输出，用户感知更快
2. 所有 LLM 调用加上 request_timeout 保护
3. /modify-document 接口也加上超时保护
"""
import os
import shutil
import json
import tempfile
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.agent.core import chat, chat_stream
from app.rag.document import index_document, search_documents, list_indexed_documents, read_document_content
from app.memory.manager import (
    get_history_messages, clear_session_history,
    create_chat, list_chats, delete_chat, rename_chat, update_chat_time,
)
from app.auth.user_manager import login_user, register_user
from app.config import settings

router = APIRouter()


# ===== 请求/响应模型 =====
class ChatRequest(BaseModel):
    """聊天请求"""
    message: str
    session_id: str = "default"


class ChatResponse(BaseModel):
    """聊天响应"""
    response: str
    session_id: str


class SearchRequest(BaseModel):
    """文档搜索请求"""
    query: str
    top_k: int = 3


class AuthRequest(BaseModel):
    """认证请求"""
    username: str
    password: str


class RenameChatRequest(BaseModel):
    """重命名会话请求"""
    username: str
    chat_id: str
    new_title: str


# ===== 认证接口 =====

@router.post("/auth/login", summary="用户登录")
async def auth_login(req: AuthRequest):
    """用户登录验证"""
    result = login_user(req.username, req.password)
    return result


@router.post("/auth/register", summary="用户注册")
async def auth_register(req: AuthRequest):
    """用户注册"""
    result = register_user(req.username, req.password)
    return result


# ===== 会话管理接口 =====

@router.post("/chats", summary="创建新会话")
async def create_new_chat(username: str = Query(..., description="用户名"), title: str = Query("新对话", description="会话标题")):
    """为用户创建一个新的聊天会话"""
    chat_info = create_chat(username, title)
    return {"success": True, "chat": chat_info}


@router.get("/chats", summary="获取用户的所有会话")
async def get_user_chats(username: str = Query(..., description="用户名")):
    """获取用户的所有聊天会话列表"""
    chats = list_chats(username)
    return {"success": True, "chats": chats}


@router.delete("/chats/{chat_id}", summary="删除会话")
async def delete_user_chat(chat_id: str, username: str = Query(..., description="用户名")):
    """删除指定的聊天会话"""
    delete_chat(username, chat_id)
    return {"success": True, "message": "会话已删除"}


@router.put("/chats/{chat_id}/rename", summary="重命名会话")
async def rename_user_chat(chat_id: str, req: RenameChatRequest):
    """重命名指定的聊天会话"""
    rename_chat(req.username, req.chat_id, req.new_title)
    return {"success": True, "message": "会话已重命名"}


# ===== 核心 Agent 接口 =====

@router.post("/chat", response_model=ChatResponse, summary="与 Agent 对话")
async def chat_api(req: ChatRequest):
    """
    核心接口：与文档助手 Agent 对话（同步版本）

    - 支持 RAG 文档问答
    - 支持员工信息查询
    - 支持多轮对话
    """
    try:
        response = chat(req.message, req.session_id)
        # 更新会话时间
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            username = parts[0]
            try:
                update_chat_time(username, req.session_id)
            except Exception:
                pass
        return ChatResponse(response=response, session_id=req.session_id)
    except Exception as e:
        error_msg = str(e)
        if "ReadTimeout" in error_msg or "timed out" in error_msg:
            raise HTTPException(status_code=504, detail="LLM 响应超时，请稍后重试")
        raise HTTPException(status_code=500, detail=f"Agent 处理失败: {error_msg}")


@router.post("/chat/stream", summary="与 Agent 流式对话")
async def chat_stream_api(req: ChatRequest):
    """
    流式对话接口（SSE - Server-Sent Events）

    实时输出 Agent 的回复，用户无需等待全部完成。
    返回格式：text/event-stream，每行 data: 后跟 JSON 片段。

    前端使用示例：
    const evtSource = new EventSource('/api/v1/chat/stream', {{
        method: 'POST',
        body: JSON.stringify({{ message: '你好', session_id: 'test' }})
    }});
    // 或用 fetch + ReadableStream
    """
    async def event_generator():
        try:
            async for chunk in chat_stream(req.message, req.session_id):
                # SSE 格式：data: {json}\n\n
                data = json.dumps({"content": chunk}, ensure_ascii=False)
                yield f"data: {data}\n\n"
            # 结束标记
            yield f"data: {json.dumps({'done': True}, ensure_ascii=False)}\n\n"
        except Exception as e:
            error_data = json.dumps({"error": str(e)}, ensure_ascii=False)
            yield f"data: {error_data}\n\n"

        # 更新会话时间
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            username = parts[0]
            try:
                update_chat_time(username, req.session_id)
            except Exception:
                pass

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # 禁止 nginx 缓冲
        },
    )


@router.post("/chat-with-file", summary="带文件的对话")
async def chat_with_file(
    file: UploadFile = File(...),
    message: str = Form(...),
    session_id: str = Form("default"),
):
    """
    带文件的对话接口（ChatGPT 风格）

    - 自动提取文件内容作为 Agent 上下文
    - 同时将文件索引到知识库
    - Agent 根据用户意图自动判断：提问则文字回答，修改则返回文件
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx", ".xlsx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}",
        )

    # 1. 保存文件到知识库目录
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 2. 索引到知识库（供后续 RAG 检索）
        try:
            index_document(file_path, file.filename)
        except Exception:
            pass  # 索引失败不影响对话

        # 3. 提取文件内容
        try:
            file_content = read_document_content(file_path)
        except Exception:
            file_content = "（无法读取文件内容）"

        # 4. 限制内容长度，避免超出 LLM 上下文
        max_content_chars = 8000
        if len(file_content) > max_content_chars:
            file_content = file_content[:max_content_chars] + f"\n\n...（文件内容过长，已截断，共 {len(file_content)} 字符）"

        # 5. 构建增强消息
        enhanced_message = f"""[用户上传了文件: {file.filename}]

文件内容如下：
---
{file_content}
---

文件保存路径: {file_path}

用户的问题/要求: {message}"""

        # 6. 调用 Agent 对话
        response = chat(enhanced_message, session_id)

        # 7. 更新会话时间
        parts = session_id.rsplit("_", 1)
        if len(parts) == 2:
            username = parts[0]
            try:
                update_chat_time(username, session_id)
            except Exception:
                pass

        return {"response": response, "session_id": session_id}

    except Exception as e:
        error_msg = str(e)
        if "ReadTimeout" in error_msg or "timed out" in error_msg:
            raise HTTPException(status_code=504, detail="LLM 响应超时，请稍后重试")
        raise HTTPException(status_code=500, detail=f"处理失败: {error_msg}")


# ===== 文档管理接口 =====

@router.post("/upload", summary="上传文档到知识库")
async def upload_document(file: UploadFile = File(...)):
    """
    上传文档并自动索引到向量数据库
    支持 PDF、TXT、DOCX、XLSX 格式
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx", ".xlsx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}",
        )

    # 保存文件
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 索引文档
    try:
        result = index_document(file_path, file.filename)
        return {"status": "success", "detail": result}
    except Exception as e:
        # 索引失败则删除文件
        os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")


@router.post("/modify-document", summary="修改文档")
async def modify_document(
    file: UploadFile = File(...),
    instruction: str = Form(...),
    username: str = Form("default"),
):
    """
    上传文档并根据修改要求进行修改，返回修改后的文件下载链接

    - 支持 PDF、TXT、DOCX、XLSX 格式
    - 修改要求：用自然语言描述如何修改
    - 返回修改后的文件下载链接
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx", ".xlsx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}",
        )

    # 保存上传的文件到临时目录
    temp_dir = os.path.join(settings.DATA_DIR, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 读取文档内容
        content = read_document_content(temp_path)

        # 调用 LLM 修改文档（带超时保护）
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
            request_timeout=settings.LLM_TIMEOUT,  # 超时保护
            max_retries=2,
        )

        # 根据文件类型调整提示词
        if ext == ".xlsx":
            system_prompt = """你是表格修改助手。按修改要求修改表格数据，只返回修改后的完整内容，不添加解释。\n\n输出格式规则（重要！）：\n- 保持 Markdown 表格格式：| 列1 | 列2 | 列3 |\n- 第一行必须是表头，第二行是 |------|------|------| 分隔线\n- 每个 Sheet 用 === Sheet: 工作表名 === 标记\n- 不要添加额外的解释或注释\n- 用中文输出"""
        else:
            system_prompt = """你是文档修改助手。按修改要求修改文档，只返回修改后的完整内容，不添加解释。保持原格式和结构，表格数据保持表格格式。用中文输出。"""

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"原始文档：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        # 保存修改后的文件
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)

        output_filename = f"modified_{file.filename}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".xlsx":
            # XLSX 格式：解析 LLM 输出的表格文本，写入 XLSX 文件
            from app.utils.xlsx_handler import write_xlsx_from_text
            actual_path = write_xlsx_from_text(modified_content, output_path, source_file=temp_path)
            output_filename = os.path.basename(actual_path)
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(temp_path)
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                paragraphs = modified_content.split("\n")
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_filename = f"modified_{os.path.splitext(file.filename)[0]}.txt"
                output_path = os.path.join(modified_dir, output_filename)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        elif ext == ".pdf":
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path, title=f"修改后的 {file.filename}")
            output_filename = os.path.basename(actual_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        # 清理临时文件
        os.remove(temp_path)

        download_url = f"/static/modified/{output_filename}"

        return {
            "success": True,
            "message": "文档修改完成！已按您的修改要求处理。",
            "download_url": download_url,
            "filename": output_filename,
        }

    except Exception as e:
        # 清理临时文件
        if os.path.exists(temp_path):
            os.remove(temp_path)
        error_msg = str(e)
        if "ReadTimeout" in error_msg or "timed out" in error_msg:
            raise HTTPException(status_code=504, detail="LLM 响应超时，请稍后重试")
        raise HTTPException(status_code=500, detail=f"文档修改失败: {error_msg}")


@router.post("/search", summary="搜索文档内容")
async def search_api(req: SearchRequest):
    """在文档库中搜索相关内容"""
    results = search_documents(req.query, req.top_k)
    return {"query": req.query, "results": results}


@router.get("/documents", summary="列出所有已索引文档")
async def list_documents():
    """获取知识库中所有文档列表"""
    docs = list_indexed_documents()
    return {"documents": docs, "count": len(docs)}


# ===== 对话历史接口 =====

@router.get("/history/{session_id}", summary="获取对话历史")
async def get_history(session_id: str):
    """获取指定会话的对话历史"""
    messages = get_history_messages(session_id)
    return {"session_id": session_id, "messages": messages, "count": len(messages)}


@router.delete("/history/{session_id}", summary="清除对话历史")
async def delete_history(session_id: str):
    """清除指定会话的对话历史"""
    clear_session_history(session_id)
    return {"status": "success", "message": f"会话 {session_id} 的历史已清除"}
