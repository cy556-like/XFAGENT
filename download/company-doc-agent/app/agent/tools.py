"""
Agent 工具定义模块
每个工具 = Agent 的一个「能力」
Agent 会根据用户问题自动选择调用哪个工具

优化策略：
1. 搜索缓存 — 相同查询 60s 内直接返回缓存结果
2. 结果截断 — 限制单条内容长度，防止上下文膨胀
3. top_k 调优 — 从 3 提升到 5，减少重复搜索
4. LLM 超时 — modify_document 工具也加上超时保护
"""
import json
import os
import time
from typing import Optional

from langchain_core.tools import tool

from app.config import settings
from app.rag.document import search_documents, index_document, list_indexed_documents


# ===== 搜索结果缓存 =====
_search_cache: dict[str, tuple[str, float]] = {}
_CACHE_TTL = 60  # 缓存有效期 60 秒

# 搜索结果单条内容最大长度（超过截断，防止上下文膨胀）
MAX_CONTENT_LENGTH = 800


def _get_cached(key: str) -> Optional[str]:
    """获取缓存"""
    if key in _search_cache:
        result, timestamp = _search_cache[key]
        if time.time() - timestamp < _CACHE_TTL:
            return result
        del _search_cache[key]
    return None


def _set_cached(key: str, value: str):
    """设置缓存"""
    _search_cache[key] = (value, time.time())
    # 清理过期缓存（简单策略：超过 20 条时清理最旧的）
    if len(_search_cache) > 20:
        oldest_key = min(_search_cache, key=lambda k: _search_cache[k][1])
        del _search_cache[oldest_key]


@tool
def search_documents_tool(query: str) -> str:
    """搜索公司文档。问制度、流程、规范时用。

    Args:
        query: 搜索查询内容
    """
    # 检查缓存
    cached = _get_cached(query)
    if cached:
        return cached

    results = search_documents(query, top_k=5)

    if not results:
        return "未找到相关文档内容。"

    output = "检索到以下相关内容：\n\n"
    for i, r in enumerate(results, 1):
        content = r['content']
        # 截断过长内容，减少上下文 token
        if len(content) > MAX_CONTENT_LENGTH:
            content = content[:MAX_CONTENT_LENGTH] + "...[内容过长已截断]"
        output += f"【文档{i}】来源: {r['source']}\n{content}\n\n"

    # 缓存结果
    _set_cached(query, output)
    return output


@tool
def lookup_employee_tool(name: str = "", department: str = "") -> str:
    """查员工信息。问部门、职位、联系方式时用。

    Args:
        name: 员工姓名（可选，模糊匹配）
        department: 部门名称（可选）
    """
    employees_file = settings.EMPLOYEES_FILE

    if not os.path.exists(employees_file):
        return "员工数据库暂未初始化，请先运行 scripts/seed_data.py 初始化数据。"

    with open(employees_file, "r", encoding="utf-8") as f:
        employees = json.load(f)

    results = employees

    # 按姓名过滤
    if name:
        results = [e for e in results if name in e.get("name", "")]

    # 按部门过滤
    if department:
        results = [e for e in results if department in e.get("department", "")]

    if not results:
        return f"未找到匹配的员工信息。（搜索条件：姓名={name}, 部门={department}）"

    output = f"找到 {len(results)} 位员工：\n\n"
    for e in results:
        output += f"{e['name']} | 部门: {e['department']} | 职位: {e['position']} | 邮箱: {e['email']}\n"
        if e.get("phone"):
            output += f"   电话: {e['phone']}\n"

    return output


@tool
def list_documents_tool() -> str:
    """列出知识库文档。"""
    docs = list_indexed_documents()

    if not docs:
        return "知识库中暂无文档。请先上传文档。"

    output = f"知识库中共有 {len(docs)} 个文档：\n\n"
    for i, doc in enumerate(docs, 1):
        output += f"  {i}. {doc}\n"

    return output


@tool
def upload_document_tool(file_path: str) -> str:
    """上传新文档到知识库。

    Args:
        file_path: 文档文件路径
    """
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"

    try:
        result = index_document(file_path)
        return f"文档上传成功！{result['message']}"
    except Exception as e:
        return f"文档上传失败: {str(e)}"


@tool
def modify_document_tool(file_path: str, instruction: str) -> str:
    """修改文档内容。用户明确要求修改/返回文件时用。

    Args:
        file_path: 要修改的文档路径
        instruction: 修改要求（自然语言描述如何修改）
    """
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"

    try:
        from app.rag.document import read_document_content
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        # 读取文档内容
        content = read_document_content(file_path)

        # 获取文件扩展名（提前获取，用于后续判断）
        ext = os.path.splitext(file_path)[1].lower()

        # 调用 LLM 修改文档（带超时保护）
        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
            request_timeout=settings.LLM_TIMEOUT,  # 超时保护
            max_retries=2,
        )

        # 根据文件类型调整提示词
        if ext == ".xlsx":
            system_prompt = """你是表格修改助手。按修改要求修改表格数据，只返回修改后的完整内容，不添加解释。

输出格式规则（重要！）：
- 保持 Markdown 表格格式：| 列1 | 列2 | 列3 |
- 第一行必须是表头，第二行是 |------|------|------| 分隔线
- ⚠️ 不要使用 === Sheet: xxx === 标记拆分多个Sheet，所有内容放在一个工作表中
- 项目信息放在表格上方的单独行中（如：项目名称：XXX），不要另建Sheet
- 不要添加额外的解释或注释
- 用中文输出"""
        else:
            system_prompt = """你是文档修改助手。按修改要求修改文档，只返回修改后的完整内容，不添加解释。保持原格式和结构。用中文输出。"""

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"原始文档：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        # 保存修改后的文件
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)
        output_filename = f"modified_{os.path.basename(file_path)}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".xlsx":
            # XLSX 格式：解析 LLM 输出的表格文本，写入 XLSX 文件
            from app.utils.xlsx_handler import write_xlsx_from_text
            actual_path = write_xlsx_from_text(modified_content, output_path, source_file=file_path)
            output_filename = os.path.basename(actual_path)
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                paragraphs = modified_content.split("\n")
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_path = output_path.replace(".docx", ".txt")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        elif ext == ".pdf":
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path, title=f"修改后的 {os.path.basename(file_path)}")
            output_filename = os.path.basename(actual_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        return f"文档修改完成！下载链接: /static/modified/{output_filename}"

    except Exception as e:
        error_msg = str(e)
        if "ReadTimeout" in error_msg or "timed out" in error_msg:
            return "文档修改超时，请稍后重试或缩短文档内容。"
        return f"文档修改失败: {error_msg}"


@tool
def create_document_tool(filename: str, content_description: str, file_format: str = "xlsx") -> str:
    """从零创建新文档（支持 XLSX/DOCX）。当用户要求"生成表格文件"、"导出XLSX"、"创建文档"且不需要基于已有文件时使用。

    Args:
        filename: 文件名（不含扩展名，如"DFMEA表"）
        content_description: 文档内容描述（包含完整的表格/文档内容，Markdown表格格式）
        file_format: 文件格式，"xlsx" 或 "docx"，默认 "xlsx"
    """
    try:
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)

        # 安全文件名：去除特殊字符
        safe_name = "".join(c for c in filename if c.isalnum() or c in "._-一二三四五六七八九十百千万亿") or "output"

        if file_format == "xlsx":
            # XLSX: 调用 LLM 生成 Markdown 表格，再解析写入 XLSX
            llm = ChatOpenAI(
                api_key=settings.LLM_API_KEY,
                base_url=settings.LLM_BASE_URL,
                model=settings.LLM_MODEL,
                temperature=0.3,
                timeout=settings.LLM_TIMEOUT,
                max_retries=1,
            )

            system_prompt = """你是表格生成助手。根据用户要求生成完整的表格数据。

输出格式规则（必须严格遵守！）：
- 使用 Markdown 表格格式：| 列1 | 列2 | 列3 |
- 第一行必须是表头，第二行是 |------|------|------| 分隔线
- ⚠️ 不要使用 === Sheet: xxx === 标记拆分多个Sheet，所有内容放在一个工作表中
- 项目信息放在表格上方的单独行中（如：项目名称：XXX），不要另建Sheet
- 严重度(S)/频度(O)/探测度(D)评级标准、AP矩阵等参考内容不需要单独建Sheet，直接省略
- 只输出表格数据，不要添加任何解释说明
- 用中文输出
- 内容要完整详细，不要省略行"""

            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=f"请生成以下表格：\n\n{content_description}"),
            ]

            response = llm.invoke(messages)
            table_content = response.content

            output_filename = f"{safe_name}.xlsx"
            output_path = os.path.join(modified_dir, output_filename)

            from app.utils.xlsx_handler import write_xlsx_from_text
            actual_path = write_xlsx_from_text(table_content, output_path)
            output_filename = os.path.basename(actual_path)

            return f"XLSX文件已生成！下载链接: /static/modified/{output_filename}"

        elif file_format == "docx":
            # DOCX: 调用 LLM 生成内容，写入 Word 文档
            llm = ChatOpenAI(
                api_key=settings.LLM_API_KEY,
                base_url=settings.LLM_BASE_URL,
                model=settings.LLM_MODEL,
                temperature=0.3,
                timeout=settings.LLM_TIMEOUT,
                max_retries=1,
            )

            system_prompt = """你是文档生成助手。根据用户要求生成完整的文档内容。
只输出文档正文，不要添加解释说明。用中文输出。"""

            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=f"请生成以下文档：\n\n{content_description}"),
            ]

            response = llm.invoke(messages)
            doc_content = response.content

            output_filename = f"{safe_name}.docx"
            output_path = os.path.join(modified_dir, output_filename)

            try:
                from docx import Document
                doc = Document()
                paragraphs = doc_content.split("\n")
                for p_text in paragraphs:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_filename = f"{safe_name}.txt"
                output_path = os.path.join(modified_dir, output_filename)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(doc_content)

            return f"DOCX文件已生成！下载链接: /static/modified/{output_filename}"

        else:
            return f"不支持的文件格式: {file_format}，仅支持 xlsx 或 docx"

    except Exception as e:
        error_msg = str(e)
        if "ReadTimeout" in error_msg or "timed out" in error_msg:
            return "文档生成超时，请稍后重试或简化内容。"
        return f"文档生成失败: {error_msg}"


# ===== 导出所有工具列表 =====
ALL_TOOLS = [
    search_documents_tool,
    lookup_employee_tool,
    list_documents_tool,
    upload_document_tool,
    modify_document_tool,
    create_document_tool,
]
