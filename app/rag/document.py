"""
文档处理与向量化模块 (RAG)
负责：加载文档 → 分块 → 向量化 → 存入 ChromaDB → 检索

优化:
- [#9] RAG 检索质量提升：混合检索（向量 + BM25关键词） + 重排序
- [#10] 引用溯源：返回结果标注文档名 + 段落位置 + chunk_id
- [#11] Embedding 降级：当 Embedding API 不可用时（403/余额不足/网络错误），
       自动切换为关键词索引模式，文档直接保存到磁盘并支持关键词搜索
- [#12] BM25 索引优化：使用 rank_bm25 库替代 Python 层全量遍历，
       从 O(N) 暴力搜索升级为倒排索引加速检索，支持 jieba 中文分词
"""
import os
os.environ["ANONYMIZED_TELEMETRY"] = "False"  # 关闭 chromadb 遥测，避免 posthog 兼容性警告
os.environ["CHROMA_TELEMETRY_ENABLED"] = "false"  # [v0.5+] 新版 chromadb 遥测开关
import re
from urllib.parse import unquote
import json
import hashlib
import logging
from typing import Optional
import time

from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
try:
    from langchain_chroma import Chroma
except ImportError:
    from langchain_community.vectorstores import Chroma

try:
    from rank_bm25 import BM25Okapi
    _RANK_BM25_AVAILABLE = True
except ImportError:
    _RANK_BM25_AVAILABLE = False
    logger_tmp = logging.getLogger(__name__)
    logger_tmp.warning("rank_bm25 未安装，BM25检索将使用旧版全量遍历模式。建议: pip install rank_bm25")

try:
    import jieba
    _JIEBA_AVAILABLE = True
except ImportError:
    _JIEBA_AVAILABLE = False

from app.config import settings

logger = logging.getLogger(__name__)

# 本地 Embedding 批量大小（本地模型无API限制，可适当增大）
EMBEDDING_BATCH_SIZE = 50  # 智谱 embedding-3 API 单次最多64条，留余量用50

# ===== 单例模式：复用 Embedding 和 ChromaDB 连接 =====
_embeddings_instance = None
_vector_store_cache = {}  # agent_id -> ChromaDB instance（按智能体隔离）

# [性能修复] 向量存储缓存上限，避免长时间运行后内存无限增长
_VECTOR_STORE_CACHE_MAX_SIZE = 20

# [优化3] ChromaDB PersistentClient 全局单例，避免重复初始化 + SQLite 锁冲突
# 每次 new PersistentClient() 会打开新的 SQLite 连接，多个实例会导致 SQLite 锁竞争
_chroma_client = None
_CHROMA_CLIENT_TTL = 1800  # [v8] 30分钟后重建，避免 SQLite WAL 锁/连接过期
# ===== [#11] Embedding 可用性标志 =====
# None = 尚未检测，True = 可用，False = 不可用（自动降级为关键词模式）
_embedding_available = None

# [优化5] Embedding 降级时间戳：记录 _embedding_available 被设为 False 的时间
# 降级 5 分钟后自动重试，不再因一次超时永久降级
_embedding_degraded_at = None
_EMBEDDING_RECOVERY_TIMEOUT = 300  # 5 分钟后自动尝试恢复

# [BUG FIX v8] Embedding 实例 TTL：空闲后 httpx 连接池中的 TCP 连接被服务端关闭
# 超时后强制重建 OpenAIEmbeddings，避免 5-30s 连接超时
_EMBEDDINGS_TTL = 900  # 15分钟

# 全局知识库的 collection 名称
GLOBAL_COLLECTION_NAME = "langchain"

# Embedding 提供者配置：
# 优先级：1. 智谱云端API（embedding-3）→ 2. 纯关键词索引
# 云端不可用时直接降级为关键词模式（本地 HuggingFace 在国内网络下无法使用，已移除）
EMBEDDING_PROVIDER = os.environ.get("EMBEDDING_PROVIDER", "openai")

# ===== [#11] 关键词索引配置 =====
KEYWORD_INDEX_DIR = os.path.join(os.path.dirname(settings.CHROMA_DIR) if hasattr(settings, 'CHROMA_DIR') else os.path.join(settings.DATA_DIR, 'keyword_index'), 'keyword_index') if hasattr(settings, 'DATA_DIR') else os.path.join(os.path.dirname(settings.CHROMA_DIR), 'keyword_index')

# 中文停用词
_STOPWORDS = {'的', '了', '是', '在', '和', '与', '有', '什么', '怎么', '如何', '哪些', '这个', '那个',
              '一个', '不是', '没有', '可以', '就是', '已经', '我们', '他们', '她们', '它们',
              '但是', '而且', '或者', '因为', '所以', '如果', '虽然', '而且', '以及',
              'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
              'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
              'should', 'may', 'might', 'shall', 'can', 'need', 'dare', 'ought',
              'used', 'to', 'of', 'in', 'for', 'on', 'with', 'at', 'by', 'from'}


def get_indexing_mode() -> str:
    """获取当前索引模式（供外部查询，如健康检查）

    Returns:
        str: "vector"（向量模式）、"keyword"（关键词模式）、"unknown"（尚未检测）
    """
    global _embedding_available
    if _embedding_available is None:
        return "unknown"
    return "vector" if _embedding_available else "keyword"


def _get_keyword_index_path(agent_id: str = None) -> str:
    """获取关键词索引JSON文件路径

    Args:
        agent_id: 智能体ID，为None时使用全局索引

    Returns:
        str: JSON文件路径
    """
    os.makedirs(KEYWORD_INDEX_DIR, exist_ok=True)
    cache_key = agent_id or "__global__"
    safe_key = cache_key.replace('-', '_').replace(' ', '_')
    return os.path.join(KEYWORD_INDEX_DIR, f"index_{safe_key}.json")


def _load_keyword_index(agent_id: str = None) -> list[dict]:
    """从磁盘加载关键词索引

    Returns:
        list[dict]: 索引条目列表，每条包含 content, source_file, chunk_index
    """
    index_path = _get_keyword_index_path(agent_id)
    if not os.path.exists(index_path):
        return []
    try:
        with open(index_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.warning(f"加载关键词索引失败: {e}")
        return []


def _save_keyword_index(index_data: list[dict], agent_id: str = None):
    """保存关键词索引到磁盘

    Args:
        index_data: 索引条目列表
        agent_id: 智能体ID
    """
    index_path = _get_keyword_index_path(agent_id)
    os.makedirs(os.path.dirname(index_path), exist_ok=True)
    try:
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"保存关键词索引失败: {e}")


def _add_chunks_to_keyword_index(chunks: list, filename: str, agent_id: str = None):
    """将文档分块添加到关键词索引

    Args:
        chunks: LangChain Document 分块列表
        filename: 文档文件名
        agent_id: 智能体ID
    """
    index_data = _load_keyword_index(agent_id)

    # 先删除该文档的旧条目（避免重复）
    index_data = [entry for entry in index_data if entry.get("source_file") != filename]

    # 添加新条目
    for chunk in chunks:
        index_data.append({
            "content": chunk.page_content,
            "source_file": filename,
            "chunk_index": chunk.metadata.get("chunk_index", 0),
        })

    _save_keyword_index(index_data, agent_id)
    # [性能优化 1] BM25 缓存失效：索引变更后清除缓存
    cache_key = agent_id or "__global__"
    if cache_key in _keyword_bm25_cache:
        del _keyword_bm25_cache[cache_key]
        logger.debug(f"[性能优化 1] BM25关键词索引缓存已失效: {cache_key}")
    logger.info(f"关键词索引已更新: {filename}, 新增 {len(chunks)} 个分块, 索引总条目={len(index_data)}")


def _delete_from_keyword_index(filename: str, agent_id: str = None) -> int:
    """从关键词索引中删除指定文档的所有条目

    Args:
        filename: 文档文件名
        agent_id: 智能体ID

    Returns:
        int: 被删除的条目数
    """
    index_data = _load_keyword_index(agent_id)
    original_count = len(index_data)
    index_data = [entry for entry in index_data if entry.get("source_file") != filename]
    deleted_count = original_count - len(index_data)

    if deleted_count > 0:
        _save_keyword_index(index_data, agent_id)
        # [性能优化 1] BM25 缓存失效：索引变更后清除缓存
        cache_key = agent_id or "__global__"
        if cache_key in _keyword_bm25_cache:
            del _keyword_bm25_cache[cache_key]
            logger.debug(f"[性能优化 1] BM25关键词索引缓存已失效: {cache_key}")
        logger.info(f"关键词索引已删除: {filename}, 删除 {deleted_count} 个条目")

    return deleted_count


def _search_keyword_index(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """[#11] 纯关键词检索（在关键词索引上搜索）
    [#12] 优化：使用 rank_bm25 + jieba 分词替代简单正则匹配
    [性能优化 1] BM25 索引缓存：避免每次搜索都重建 BM25Okapi（节省 200-800ms/次）

    不依赖 ChromaDB/Embedding，直接在 JSON 索引上做关键词匹配。
    当 Embedding API 不可用时作为主要检索手段。

    Args:
        query: 用户查询
        top_k: 返回最相关的K个结果
        agent_id: 智能体ID

    Returns:
        list[dict]: 检索结果列表
    """
    index_data = _load_keyword_index(agent_id)
    if not index_data:
        return []

    # ===== 计算 index hash 用于缓存失效检测 =====
    cache_key = agent_id or "__global__"
    now = time.time()
    # 使用索引条目数和首尾内容的 hash 快速检测变化
    index_hash = hashlib.md5(
        (str(len(index_data)) + 
         (index_data[0].get("content", "") if index_data else "") +
         (index_data[-1].get("content", "") if index_data else "")
        ).encode()
    ).hexdigest()

    # ===== [#12] 优先使用 rank_bm25 索引（带缓存） =====
    if _RANK_BM25_AVAILABLE:
        try:
            # [性能优化 1] 检查缓存：相同 agent_id + index_hash 复用 BM25 索引
            cached = _keyword_bm25_cache.get(cache_key)
            if cached is not None and cached.get("index_hash") == index_hash and \
               now - cached.get("updated_at", 0) < _KEYWORD_BM25_TTL:
                bm25 = cached["bm25"]
                corpus = cached["corpus"]
                cached_index_data = cached["index_data"]
            else:
                # 缓存失效，重建 BM25 索引
                corpus = [entry.get("content", "") for entry in index_data if entry.get("content")]
                if not corpus:
                    return []

                tokenized_corpus = [_tokenize_text(doc) for doc in corpus]
                bm25 = BM25Okapi(tokenized_corpus)

                # LRU 淘汰
                while len(_keyword_bm25_cache) >= _KEYWORD_BM25_MAX_ENTRIES:
                    oldest_key = min(_keyword_bm25_cache, key=lambda k: _keyword_bm25_cache[k].get("updated_at", 0))
                    del _keyword_bm25_cache[oldest_key]

                _keyword_bm25_cache[cache_key] = {
                    "bm25": bm25,
                    "corpus": corpus,
                    "index_data": index_data,
                    "updated_at": now,
                    "index_hash": index_hash,
                }
                cached_index_data = index_data
                logger.debug(f"[性能优化 1] BM25关键词索引已构建并缓存: cache_key={cache_key}, 条目={len(corpus)}")

            tokenized_query = _tokenize_text(query)
            if not tokenized_query:
                return []

            doc_scores = bm25.get_scores(tokenized_query)

            # 收集得分 > 0 的结果
            scored = []
            for i, score in enumerate(doc_scores):
                if score <= 0:
                    continue
                entry = cached_index_data[i]
                scored.append({
                    "content": entry.get("content", ""),
                    "source": entry.get("source_file", "未知来源"),
                    "chunk_index": entry.get("chunk_index", -1),
                    "relevance_score": round(float(score), 4),
                })

            scored.sort(key=lambda x: x["relevance_score"], reverse=True)
            return scored[:top_k]
        except Exception as e:
            logger.warning(f"[#12] 关键词索引 rank_bm25 搜索失败，降级为正则匹配: {e}")

    # ===== 降级：旧版正则匹配 =====
    query_terms = set(re.findall(r'[\u4e00-\u9fff]+|\w+', query.lower()))
    query_terms = query_terms - _STOPWORDS

    if not query_terms:
        query_terms = {query.lower()}

    scored = []
    for entry in index_data:
        content = entry.get("content", "")
        content_lower = content.lower()

        match_count = sum(1 for term in query_terms if term in content_lower)
        if match_count == 0:
            continue

        term_coverage = match_count / max(len(query_terms), 1)
        tf_score = match_count / max(len(content), 1) * 1000
        combined_score = tf_score * 0.6 + term_coverage * 100 * 0.4

        scored.append({
            "content": content,
            "source": entry.get("source_file", "未知来源"),
            "chunk_index": entry.get("chunk_index", -1),
            "relevance_score": round(combined_score, 4),
        })

    scored.sort(key=lambda x: x["relevance_score"], reverse=True)
    return scored[:top_k]


def _is_embedding_error(e: Exception) -> bool:
    """判断异常是否为 Embedding 不可用错误

    [优化6] 错误判断收窄：超时/连接错误不再标记为永久不可用。
    只有明确的认证/权限/余额错误才标记为永久降级，
    超时和临时网络错误不应导致永久降级（可能只是瞬时波动）。
    
    永久降级（返回 True）：
    - 403 权限/认证错误
    - 429 余额/限流错误
    - 1113 智谱余额不足
    - API key 错误
    - new_api_error（智谱特有）
    
    临时错误（返回 False，由 [优化5] 自动恢复机制处理）：
    - timeout / connection / connect 等网络超时错误
    """
    error_str = str(e).lower()
    # 403 权限/认证错误 → 永久降级
    if '403' in error_str or 'no access' in error_str or 'forbidden' in error_str:
        return True
    # 429 余额/限流错误 → 永久降级
    if '429' in error_str or '余额' in error_str or 'rate limit' in error_str or 'quota' in error_str:
        return True
    # 1113 智谱余额不足 → 永久降级
    if '1113' in error_str:
        return True
    # API key 错误 → 永久降级
    if 'api_key' in error_str or 'api key' in error_str or 'unauthorized' in error_str or 'invalid api' in error_str:
        return True
    # new_api_error（智谱特有）→ 永久降级
    if 'new_api_error' in error_str:
        return True
    # [优化6] 超时/连接错误 → 不再标记为永久不可用，由自动恢复机制处理
    # 之前的逻辑把 timeout/connection 也标记为永久降级，导致一次网络波动就永久降级
    # 现在这些临时错误返回 False，配合 _embedding_degraded_at + _EMBEDDING_RECOVERY_TIMEOUT 实现自动恢复
    return False


def get_embeddings():
    """获取 Embedding 模型（单例复用，避免重复初始化）

    优先级（两级自动降级）：
    1. 智谱云端 API Embedding（embedding-3）— 质量好，需联网+额度
    2. 纯关键词索引 — 无需任何模型，兜底方案

    注意：本地 HuggingFace Embedding 已移除（国内网络无法连接 huggingface.co）

    [#11] 当 Embedding 不可用时，标记 _embedding_available = False，
    后续操作将自动降级为关键词索引模式
    
    [优化5] Embedding 自动恢复：降级 5 分钟后自动重试。
    之前一次超时就永久降级，现在 _embedding_degraded_at 记录降级时间，
    超过 _EMBEDDING_RECOVERY_TIMEOUT（300秒）后重新尝试初始化。
    """
    global _embeddings_instance, _embedding_available, _embedding_degraded_at

    # 如果已知 Embedding 不可用，检查是否已超过恢复时间
    if _embedding_available is False:
        if _embedding_degraded_at is not None:
            elapsed = time.time() - _embedding_degraded_at
            if elapsed >= _EMBEDDING_RECOVERY_TIMEOUT:
                # [优化5] 超过恢复时间，自动重试
                logger.info(f"[优化5] Embedding 降级已 {elapsed:.0f}s，超过恢复阈值 {_EMBEDDING_RECOVERY_TIMEOUT}s，自动重试")
                _embedding_available = None  # 重置为未检测状态，让下面的逻辑重新初始化
                _embeddings_instance = None
                _embedding_degraded_at = None
            else:
                logger.debug(f"Embedding 降级中，剩余 { _EMBEDDING_RECOVERY_TIMEOUT - elapsed:.0f}s 后自动重试")
                return None
        else:
            return None

    # [BUG FIX v8] TTL 检查：空闲后重建 Embedding 实例，避免死 TCP 连接
    if (_embeddings_instance is not None
        and hasattr(_embeddings_instance, '_created_at')
        and time.time() - _embeddings_instance._created_at > _EMBEDDINGS_TTL):
        logger.info(f"Embedding 实例超过 TTL（{_EMBEDDINGS_TTL}s），重建以刷新连接池")
        _embeddings_instance = None

    if _embeddings_instance is None:
        try:
            from langchain_openai import OpenAIEmbeddings
            embedding_model = getattr(settings, 'EMBEDDING_MODEL', 'embedding-3')
            # [#12] Embedding 使用独立 API Key（EMBEDDING_API_KEY），如未设置则回退到 LLM_API_KEY
            embedding_api_key = getattr(settings, 'EMBEDDING_API_KEY', '') or settings.LLM_API_KEY
            embedding_base_url = getattr(settings, 'EMBEDDING_BASE_URL', '') or settings.LLM_BASE_URL
            _embeddings_instance = OpenAIEmbeddings(
                api_key=embedding_api_key,
                base_url=embedding_base_url,
                model=embedding_model,
                request_timeout=15,  # 超时保护：15秒，避免初始化卡死整个服务
            )
            _embeddings_instance._created_at = time.time()  # [v8] 记录创建时间用于 TTL 检查
            _embedding_available = True
            logger.info(f"✅ Embedding 模型已初始化（智谱云端API）: {embedding_model}")
        except ImportError:
            logger.error("langchain-openai 未安装，Embedding 不可用")
            _embeddings_instance = None
            _embedding_available = False
            _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
            logger.warning("❌ Embedding 不可用，系统将使用关键词索引模式")
        except Exception as e:
            logger.error(f"云端 Embedding 初始化失败: {e}")
            _embeddings_instance = None
            _embedding_available = False
            _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
            logger.warning("❌ Embedding 不可用，系统将使用关键词索引模式")

    return _embeddings_instance


def _get_collection_name(agent_id: str = None) -> str:
    """根据 agent_id 获取 ChromaDB collection 名称

    - agent_id 为 None 或空 → 全局知识库（普通Agent模式）
    - agent_id 有值 → 智能体专属知识库
    """
    if agent_id:
        # 用 agent_id 做 collection 名，确保合法
        safe_id = agent_id.replace('-', '_').replace(' ', '_')
        return f"agent_{safe_id}"
    return GLOBAL_COLLECTION_NAME


def get_vector_store(agent_id: str = None):
    """获取 ChromaDB 向量数据库实例（按 agent_id 隔离）

    Args:
        agent_id: 智能体ID，为 None 时使用全局知识库

    每个智能体有独立的 ChromaDB collection，互不干扰。
    普通 Agent 模式使用默认的全局 collection。

    [#11] 当 Embedding 不可用时返回 None
    
    [优化3] 使用全局 PersistentClient 单例，避免重复初始化和 SQLite 锁冲突
    """
    global _embedding_available, _embedding_degraded_at, _chroma_client

    # 如果已知 Embedding 不可用，检查自动恢复
    if _embedding_available is False:
        # 委托 get_embeddings() 检查恢复逻辑
        embeddings = get_embeddings()
        if embeddings is None:
            return None

    cache_key = agent_id or "__global__"

    if cache_key not in _vector_store_cache:
        embeddings = get_embeddings()
        if embeddings is None:
            return None

        collection_name = _get_collection_name(agent_id)
        try:
            # [优化3] 使用全局 PersistentClient 单例，避免重复创建导致的 SQLite 锁冲突
            # [v8] TTL 检查：空闲后重建，避免 SQLite 连接过期
            if (_chroma_client is not None
                and hasattr(_chroma_client, '_created_at')
                and time.time() - _chroma_client._created_at > _CHROMA_CLIENT_TTL):
                logger.info(f"ChromaDB PersistentClient 超过 TTL（{_CHROMA_CLIENT_TTL}s），重建连接")
                _chroma_client = None
                _vector_store_cache.clear()  # 旧的 Chroma 实例引用已失效
            
            if _chroma_client is None:
                import chromadb
                _chroma_client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
                _chroma_client._created_at = time.time()  # [v8] 记录创建时间
                logger.info(f"[优化3] ChromaDB PersistentClient 已创建全局单例: {settings.CHROMA_DIR}")
            vs = Chroma(
                collection_name=collection_name,
                client=_chroma_client,
                embedding_function=embeddings,
            )
            _vector_store_cache[cache_key] = vs
            logger.info(f"ChromaDB 已连接: collection={collection_name}, agent_id={agent_id}")
            # [性能修复] LRU淘汰：超过上限时移除最早的缓存
            while len(_vector_store_cache) > _VECTOR_STORE_CACHE_MAX_SIZE:
                oldest_key = next(iter(_vector_store_cache))
                del _vector_store_cache[oldest_key]
                logger.info(f"[性能修复] 向量存储缓存淘汰: {oldest_key}")
        except Exception as e:
            logger.error(f"ChromaDB 连接失败: {e}")
            # 重试：重建 PersistentClient 单例
            try:
                import chromadb
                _chroma_client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
                vs = Chroma(
                    collection_name=collection_name,
                    client=_chroma_client,
                    embedding_function=embeddings,
                )
                _vector_store_cache[cache_key] = vs
                logger.info(f"ChromaDB 已连接(retry): collection={collection_name}")
                # [性能修复] LRU淘汰
                while len(_vector_store_cache) > _VECTOR_STORE_CACHE_MAX_SIZE:
                    oldest_key = next(iter(_vector_store_cache))
                    del _vector_store_cache[oldest_key]
            except Exception as e2:
                logger.error(f"ChromaDB 连接失败(retry): {e2}")
                if _is_embedding_error(e) or _is_embedding_error(e2):
                    _embedding_available = False
                    _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
                    logger.warning(f"Embedding 不可用，已记录降级时间，{_EMBEDDING_RECOVERY_TIMEOUT}s 后自动重试")
                return None

    return _vector_store_cache.get(cache_key)


def reset_vector_store():
    """重置向量数据库单例（配置变更时调用）"""
    global _embeddings_instance, _embedding_available, _embedding_degraded_at, _chroma_client
    _vector_store_cache.clear()
    _embeddings_instance = None
    _embedding_available = None  # 重置后重新检测
    _embedding_degraded_at = None  # [优化5] 清除降级时间
    _chroma_client = None  # [优化3] 清除单例，让下次重新创建
    logger.info("向量数据库单例已重置，将重新检测 Embedding 可用性")


def reindex_all_documents(agent_id: str = None):
    """重建指定知识库的所有文档索引（切换embedding模型后调用）

    当从 OpenAI Embedding 切换到本地 Embedding 时，
    旧向量数据的维度不同，需要删除旧collection并重新索引。

    [#11] 同时支持向量模式和关键词模式

    Args:
        agent_id: 智能体ID，为None时重建全局知识库

    Returns:
        dict: 包含重建结果
    """
    import chromadb
    collection_name = _get_collection_name(agent_id)

    try:
        # 1. 收集所有文档来源（ChromaDB + 关键词索引 + 磁盘文件）
        document_files = set()

        # 从 ChromaDB 获取
        try:
            client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
            existing_collections = [c.name for c in client.list_collections()]
            if collection_name in existing_collections:
                collection = client.get_collection(collection_name)
                all_docs = collection.get(include=["metadatas"])
                for meta in (all_docs.get("metadatas") or []):
                    if meta and "source_file" in meta:
                        document_files.add(meta["source_file"])
                # 删除旧collection
                client.delete_collection(collection_name)
                logger.info(f"已删除旧collection: {collection_name}")
        except Exception as e:
            logger.warning(f"从ChromaDB获取文档列表失败: {e}")

        # 从关键词索引获取
        keyword_docs = _load_keyword_index(agent_id)
        for entry in keyword_docs:
            if entry.get("source_file"):
                document_files.add(entry["source_file"])

        # 从磁盘扫描
        if agent_id:
            scan_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")
        else:
            scan_dir = settings.DOCUMENTS_DIR
        if os.path.exists(scan_dir):
            for fname in os.listdir(scan_dir):
                ext = os.path.splitext(fname)[1].lower()
                if ext in {'.pdf', '.txt', '.docx'}:
                    file_path = os.path.join(scan_dir, fname)
                    if os.path.isfile(file_path):
                        document_files.add(fname)

        # 2. 清除缓存
        cache_key = agent_id or "__global__"
        if cache_key in _vector_store_cache:
            del _vector_store_cache[cache_key]

        # 3. 清除旧关键词索引
        keyword_index_path = _get_keyword_index_path(agent_id)
        if os.path.exists(keyword_index_path):
            os.remove(keyword_index_path)

        # 4. 重新索引所有文档
        reindexed = []
        failed = []
        for filename in document_files:
            # 查找文件路径（可能在agent子目录中）
            if agent_id:
                file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)
            else:
                file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

            if not os.path.exists(file_path):
                # 尝试全局目录
                file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

            if os.path.exists(file_path):
                try:
                    result = index_document(file_path, filename, agent_id=agent_id)
                    reindexed.append(filename)
                    logger.info(f"重新索引成功: {filename}, {result.get('chunks', 0)} 个分块")
                except Exception as e:
                    failed.append(f"{filename}: {str(e)}")
                    logger.error(f"重新索引失败: {filename}, {e}")
            else:
                failed.append(f"{filename}: 文件不存在")

        mode_str = "向量" if _embedding_available else "关键词"
        return {
            "status": "success",
            "collection": collection_name,
            "indexing_mode": mode_str,
            "documents_found": len(document_files),
            "reindexed": len(reindexed),
            "failed": failed,
            "message": f"知识库重建完成（{mode_str}模式）: 找到{len(document_files)}个文档，成功索引{len(reindexed)}个" + (f"，失败{len(failed)}个" if failed else "")
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"重建知识库失败: {str(e)}"
        }


def _load_docx_with_tables(file_path: str) -> list:
    """使用 python-docx 加载 DOCX 文件，保留表格结构为 Markdown 格式

    解决 Docx2txtLoader 的核心缺陷：
    - Docx2txtLoader 把所有表格展平为纯文本，丢失列对齐和结构信息
    - 宽表格（如28列DFMEA表）展平后完全无法理解
    - 本函数将表格转为 Markdown 表格语法，LLM 可正确理解表格结构

    策略：
    - 窄表格（≤8列）：直接转为 Markdown 表格
    - 宽表格（>8列）：按分组标题拆分为多个子表格，每个子表格展示一个分组的列
      例如28列DFMEA表拆分为：结构分析、功能分析、失效分析、风险分析、优化 五个子表
    - 2列键值表格：转为「键：值」格式，更紧凑易读

    Returns:
        list[Document]: LangChain Document 列表，每个 Document 的 page_content 包含 Markdown 格式文本
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        # python-docx 未安装，回退到 Docx2txtLoader
        logger.warning("python-docx 未安装，回退到 Docx2txtLoader（表格结构将丢失）")
        loader = Docx2txtLoader(file_path)
        return loader.load()

    doc = DocxDocument(file_path)
    content_parts = []

    # 遍历文档所有元素（段落和表格），按文档顺序输出
    # python-docx 的 doc.element.body 包含所有块级元素
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    table_index = 0  # 跟踪当前处理到第几个表格

    for element in body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # 段落元素
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            # 识别标题级别
            style_name = para.style.name if para.style else ''
            if 'Heading 1' in style_name or 'heading 1' in style_name:
                content_parts.append(f'# {text}')
            elif 'Heading 2' in style_name or 'heading 2' in style_name:
                content_parts.append(f'## {text}')
            elif 'Heading 3' in style_name or 'heading 3' in style_name:
                content_parts.append(f'### {text}')
            elif 'Heading 4' in style_name or 'heading 4' in style_name:
                content_parts.append(f'#### {text}')
            else:
                content_parts.append(text)

        elif tag == 'tbl':
            # 表格元素
            if table_index >= len(doc.tables):
                table_index += 1
                continue
            table = doc.tables[table_index]
            table_index += 1

            md_table = _convert_table_to_markdown(table)
            if md_table:
                content_parts.append(md_table)

    full_content = '\n\n'.join(content_parts)

    if not full_content.strip():
        return []

    from langchain_core.documents import Document
    return [Document(page_content=full_content, metadata={"source": file_path})]


def _convert_table_to_markdown(table) -> str:
    """将 python-docx Table 转换为 Markdown 格式

    策略：
    - 2列键值表 → 紧凑的「键：值」格式
    - 窄表格（3-8列）→ 标准 Markdown 表格
    - 宽表格（>8列）→ 按分组拆分为多个子表格

    Args:
        table: python-docx Table 对象

    Returns:
        str: Markdown 格式的表格文本
    """
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
            # 清理多余空白
            cell_text = re.sub(r'\s+', ' ', cell_text).strip()
            cells.append(cell_text)
        rows_data.append(cells)

    if not rows_data:
        return ''

    num_cols = max(len(row) for row in rows_data)
    # 补齐短行
    for row in rows_data:
        while len(row) < num_cols:
            row.append('')

    # 过滤全空行
    rows_data = [row for row in rows_data if any(cell.strip() for cell in row)]
    if not rows_data:
        return ''

    num_cols = max(len(row) for row in rows_data)

    # ===== 2列键值表 → 紧凑格式 =====
    if num_cols == 2:
        # 检查是否为键值对模式（第一列短，第二列长）
        first_col_lens = [len(row[0]) for row in rows_data if row[0].strip()]
        second_col_lens = [len(row[1]) for row in rows_data if row[1].strip()]
        avg_first = sum(first_col_lens) / max(len(first_col_lens), 1)
        avg_second = sum(second_col_lens) / max(len(second_col_lens), 1)

        if avg_first < 20 and avg_first < avg_second * 0.6:
            # 键值对模式：用「键：值」格式
            lines = []
            for row in rows_data:
                key = row[0].strip()
                val = row[1].strip()
                if key and val:
                    lines.append(f'**{key}**：{val}')
                elif val:
                    lines.append(val)
                elif key:
                    lines.append(f'**{key}**')
            return '\n'.join(lines)

    # ===== 窄表格（3-8列）→ 标准 Markdown 表格 =====
    if num_cols <= 8:
        return _format_narrow_table(rows_data, num_cols)

    # ===== 宽表格（>8列）→ 按分组拆分 =====
    return _format_wide_table(rows_data, num_cols)


def _format_narrow_table(rows_data: list, num_cols: int) -> str:
    """将窄表格（≤8列）格式化为标准 Markdown 表格"""
    lines = []

    # 表头行
    header = rows_data[0]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * num_cols) + ' |')

    # 数据行
    for row in rows_data[1:]:
        # 补齐列数
        padded = row + [''] * (num_cols - len(row))
        lines.append('| ' + ' | '.join(padded[:num_cols]) + ' |')

    return '\n'.join(lines)


def _format_wide_table(rows_data: list, num_cols: int) -> str:
    """将宽表格（>8列）转为结构化键值格式

    宽表格（如28列DFMEA表）直接作为Markdown表格无法阅读。
    策略：将每行数据按分组转为键值对格式，类似第二个DFMEA文件的结构。

    支持两种分组模式：
    1. 标准两级表头：Row 0 有分组名（如"结构分析"），Row 1 有子列名（如"上一较高级别"）
       → 按分组输出键值对
    2. 无明显分组：所有列平铺为键值对

    输出格式示例：
    【条目1】凸模组件 - 凸模本体
    结构分析：
    - 上一较高级别：冷冲压模具总成
    - 关注要素：凸模组件
    - 下一较低级别/特性：凸模本体
    功能分析：
    - 上一较高级别功能：稳定冲压出合格零件
    ...
    """
    if len(rows_data) < 2:
        return _format_narrow_table(rows_data, num_cols)

    # 分析 Row 0 和 Row 1 的关系
    first_row = rows_data[0]
    second_row = rows_data[1] if len(rows_data) > 1 else []

    # 检测 Row 0 是否为分组标题行
    # 特征：Row 0 非空单元格数量 << Row 1 非空单元格数量
    first_nonempty = sum(1 for c in first_row if c.strip())
    second_nonempty = sum(1 for c in second_row if c.strip())

    # 判断 Row 1 是否为子列标题行（分隔行或表头行）
    is_row1_separator = all(re.match(r'^[-:]+$', c.strip()) for c in second_row if c.strip())

    if is_row1_separator:
        # Row 1 是分隔行，说明 Row 0 就是唯一的表头行
        # 直接将宽表格每行转为键值对
        return _wide_table_to_kv_flat(first_row, rows_data[1:], num_cols)

    # Row 0 有分组，Row 1 有子列标题
    # 构建分组映射：group_name -> [(col_index, sub_col_name), ...]
    groups = _build_column_groups(first_row, second_row, num_cols)

    if not groups or len(groups) <= 1:
        # 无法识别有效分组，回退为键值对格式
        # 使用 Row 1 作为主列名，Row 0 作为备用（当 Row 1 为分隔线时回退）
        use_second = second_nonempty > first_nonempty
        return _wide_table_to_kv_flat(
            second_row if use_second else first_row,
            rows_data[2:] if use_second else rows_data[1:],
            num_cols,
            fallback_row=first_row  # Row 0 作为备用列名来源
        )

    # 按分组输出每行数据的键值对
    # 确定数据行起始位置（跳过 Row 0 分组标题 + Row 1 子列标题）
    data_start = 2
    seq_col = 0  # 序号列索引

    parts = []
    for row_idx in range(data_start, len(rows_data)):
        row = rows_data[row_idx]

        # 获取序号
        seq_val = row[seq_col].strip() if seq_col < len(row) else ''

        # 构建条目标题：从第一组中提取关键信息
        title_parts = []
        for group_name, col_items in groups:
            for col_idx, col_name in col_items:
                val = row[col_idx].strip() if col_idx < len(row) else ''
                if val and col_idx <= 3:  # 取前几列关键信息
                    title_parts.append(val)
            if title_parts:
                break  # 只从第一组取标题

        if seq_val:
            parts.append(f'【条目{seq_val}】{" - ".join(title_parts)}')
        elif title_parts:
            parts.append(f'【条目】{" - ".join(title_parts)}')

        # 按分组输出键值对
        for group_name, col_items in groups:
            group_lines = []
            for col_idx, col_name in col_items:
                if col_idx == seq_col:
                    continue  # 跳过序号列
                val = row[col_idx].strip() if col_idx < len(row) else ''
                if val:
                    group_lines.append(f'- {col_name}：{val}')

            if group_lines:
                parts.append(f'{group_name}：')
                parts.extend(group_lines)

        parts.append('')  # 条目间空行

    return '\n'.join(parts)


def _build_column_groups(first_row: list, second_row: list, num_cols: int) -> list:
    """构建宽表格的列分组映射

    分析 Row 0（分组标题行）和 Row 1（子列标题行），
    确定每个分组包含哪些列。

    策略（按优先级）：
    1. 如果 Row 0 有分组标题且与 Row 1 形成正确的"一对多"关系 → 使用 Row 0 分组
    2. 如果 Row 0 非空单元格只是与 Row 1 一一对应 → 无分组，用 Row 1 做列名
    3. 如果 Row 0 分组不可靠（每个分组只覆盖1列） → 退回到 Row 1 做列名

    可靠分组判断：
    - Row 0 非空单元格数量 << Row 1 非空单元格数量
    - 分组后每个分组平均覆盖 >1 列
    - 总覆盖列数 >= 总列数的50%

    Args:
        first_row: 第一行数据
        second_row: 第二行数据
        num_cols: 总列数

    Returns:
        list[tuple]: [(分组名, [(列索引, 子列名), ...]), ...]  空列表表示无法分组
    """
    first_nonempty_cells = [(i, first_row[i].strip()) for i in range(min(len(first_row), num_cols)) if first_row[i].strip()]

    # 如果 Row 0 非空单元格太少，无法分组
    if len(first_nonempty_cells) <= 1:
        return []

    # 如果 Row 0 的非空单元格数量接近总列数（>=60%），说明不是分组标题行
    if len(first_nonempty_cells) >= num_cols * 0.6:
        return []

    # 按非空单元格位置确定分组边界
    raw_groups = []
    for gi, (cell_idx, group_name) in enumerate(first_nonempty_cells):
        if gi + 1 < len(first_nonempty_cells):
            next_group_start = first_nonempty_cells[gi + 1][0]
        else:
            next_group_start = num_cols
        col_range = list(range(cell_idx, next_group_start))
        raw_groups.append((group_name, col_range))

    # 跳过序号列分组
    groups_filtered = []
    for name, cols in raw_groups:
        if len(cols) == 1 and name in ('序号', '编号', 'No', '#', '---'):
            continue
        groups_filtered.append((name, cols))

    if not groups_filtered:
        return []

    # 可靠性检查：如果大多数分组只覆盖1列，说明 Row 0 没有正确的单元格合并
    # 这种情况下分组是错误的，应退回到用 Row 1 做列名
    total_covered = sum(len(cols) for _, cols in groups_filtered)
    single_col_groups = sum(1 for _, cols in groups_filtered if len(cols) <= 1)

    # 如果超过60%的分组只有1列，或总覆盖列数不到50%，判定分组不可靠
    if single_col_groups > len(groups_filtered) * 0.6 or total_covered < num_cols * 0.5:
        logger.debug(f"宽表格分组不可靠: {single_col_groups}/{len(groups_filtered)} 组只有1列, "
                     f"覆盖 {total_covered}/{num_cols} 列, 退回到键值对格式")
        return []

    # 为每列匹配子列名（来自 Row 1）
    result = []
    for group_name, col_range in groups_filtered:
        col_items = []
        for ci in col_range:
            sub_name = second_row[ci].strip() if ci < len(second_row) else ''
            if re.match(r'^[-:]+$', sub_name):
                sub_name = f'列{ci + 1}'
            if not sub_name:
                sub_name = f'列{ci + 1}'
            col_items.append((ci, sub_name))
        result.append((group_name, col_items))

    return result


def _wide_table_to_kv_flat(header_row: list, data_rows: list, num_cols: int,
                            fallback_row: list = None) -> str:
    """将宽表格转为平铺的键值对格式（无分组）

    当无法识别分组时，每行数据转为：
    - 列名1：值1
    - 列名2：值2
    ...

    Args:
        header_row: 表头行（通常是 Row 1 的子列标题）
        data_rows: 数据行列表
        num_cols: 总列数
        fallback_row: 备用标题行（如 Row 0），当 header_row 对应位置为分隔线或空时使用

    Returns:
        str: 键值对格式的文本
    """
    # 构建列名（处理分隔线、空值、重复列名）
    col_names = []
    seen_names = {}  # 列名出现次数，用于去重
    for i in range(num_cols):
        name = header_row[i].strip() if i < len(header_row) else ''
        # 分隔线或空值 → 尝试从 fallback_row 获取
        if not name or re.match(r'^[-:]+$', name):
            if fallback_row and i < len(fallback_row):
                name = fallback_row[i].strip()
        # 仍然为空 → 用列号
        if not name or re.match(r'^[-:]+$', name):
            name = f'列{i + 1}'

        # 处理重复列名：第二次出现时添加后缀
        if name in seen_names:
            seen_names[name] += 1
            # 常见重复模式：严重度/频度/探测度/AP 的"措施后"版本
            name = f'{name}（措施后）'
            # 如果仍然重复（极端情况），加序号
            if name in seen_names:
                seen_names[name] += 1
                name = f'{name}{seen_names[name]}'
            else:
                seen_names[name] = 1
        else:
            seen_names[name] = 1

        col_names.append(name)

    parts = []
    for row_idx, row in enumerate(data_rows):
        # 跳过分隔行（所有非空单元格都是分隔线 ---）
        non_empty_cells = [row[ci].strip() for ci in range(min(len(row), num_cols)) if ci < len(row) and row[ci].strip()]
        if non_empty_cells and all(re.match(r'^[-:]+$', c) for c in non_empty_cells):
            continue
        # 跳过与表头完全相同的行（重复表头）
        if all(row[ci].strip() == header_row[ci].strip() if ci < len(row) and ci < len(header_row) else True for ci in range(min(len(row), num_cols))):
            continue

        seq_val = row[0].strip() if len(row) > 0 else ''
        if seq_val:
            parts.append(f'【条目{seq_val}】')
        else:
            parts.append(f'【条目{row_idx + 1}】')

        for ci in range(num_cols):
            val = row[ci].strip() if ci < len(row) else ''
            if val:
                parts.append(f'- {col_names[ci]}：{val}')

        parts.append('')  # 条目间空行

    return '\n'.join(parts)


def load_xlsx_document(file_path: str) -> list:
    """加载 XLSX/XLS 文件，将每个工作表的内容转为文本格式供 LLM 分析
    
    策略：
    - 每个工作表单独输出
    - 表格内容转为 Markdown 表格语法
    - 支持合并单元格（展开为多行多列）
    
    Returns:
        list[Document]: LangChain Document 列表
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.warning("openpyxl 未安装，无法读取 xlsx 文件")
        return []
    
    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        content_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            content_parts.append(f'# 工作表：{sheet_name}')
            
            rows_data = []
            for row in ws.iter_rows(values_only=True):
                cells = []
                for cell in row:
                    if cell is None:
                        cells.append('')
                    else:
                        cells.append(str(cell).strip())
                # Skip completely empty rows
                if any(c for c in cells):
                    rows_data.append(cells)
            
            if not rows_data:
                content_parts.append('（空工作表）')
                continue
            
            # Determine column count
            num_cols = max(len(row) for row in rows_data)
            # Pad short rows
            for row in rows_data:
                while len(row) < num_cols:
                    row.append('')
            
            # Convert to Markdown table
            if num_cols <= 8:
                # Narrow table: standard Markdown table
                # Header row
                header = rows_data[0]
                content_parts.append('| ' + ' | '.join(header) + ' |')
                content_parts.append('| ' + ' | '.join(['---'] * num_cols) + ' |')
                # Data rows
                for row in rows_data[1:]:
                    content_parts.append('| ' + ' | '.join(row[:num_cols]) + ' |')
            else:
                # Wide table: key-value format
                header = rows_data[0]
                for row_idx, row in enumerate(rows_data[1:], 1):
                    title_val = row[0].strip() if row[0].strip() else f'行{row_idx}'
                    content_parts.append(f'【条目{row_idx}】{title_val}')
                    for ci in range(num_cols):
                        col_name = header[ci].strip() if ci < len(header) and header[ci].strip() else f'列{ci+1}'
                        val = row[ci].strip() if ci < len(row) else ''
                        if val:
                            content_parts.append(f'- {col_name}：{val}')
                    content_parts.append('')
            
            content_parts.append('')
        
        wb.close()
        full_content = '\n'.join(content_parts)
        
        if not full_content.strip():
            return []
        
        from langchain_core.documents import Document
        return [Document(page_content=full_content, metadata={"source": file_path})]
    except Exception as e:
        logger.error(f"加载 XLSX 文件失败: {e}")
        return []


def load_document(file_path: str) -> list:
    """
    根据文件类型加载文档
    支持：PDF、TXT、MD、DOCX、XLSX、XLS

    DOCX 文件使用 python-docx 加载，保留表格结构为 Markdown 格式，
    解决 Docx2txtLoader 展平表格导致结构丢失的问题。
    XLSX/XLS 文件使用 openpyxl 加载，将工作表转为 Markdown 表格格式。
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    elif ext == ".md":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    elif ext == ".docx":
        return _load_docx_with_tables(file_path)
    elif ext in (".xlsx", ".xls"):
        return load_xlsx_document(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 PDF/TXT/MD/DOCX/XLSX/XLS")


def _split_markdown_by_headers(docs: list, chunk_size: int = 800, chunk_overlap: int = 200) -> list:
    """Markdown 文件专用切片：按标题层级切分，保留标题层级信息到 metadata

    与 RecursiveCharacterTextSplitter 的区别：
    - RecursiveCharacter 只是按 # 符号作为分隔符切割，切割后的 chunk 不保留标题信息
    - MarkdownHeaderTextSplitter 按标题层级切割，每个 chunk 的 metadata 自动带上各级标题
      例如：{"Header 1": "动力电池系统", "Header 2": "故障诊断"}

    这样检索时 AI 能知道这个 chunk 属于哪个章节，回答更精准。
    """
    md_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=[
            ("#", "Header 1"),
            ("##", "Header 2"),
            ("###", "Header 3"),
            ("####", "Header 4"),
        ],
        strip_headers=False,  # 保留标题文本在 chunk 内容中
    )

    all_chunks = []
    for doc in docs:
        try:
            md_chunks = md_splitter.split_text(doc.page_content)
        except Exception as e:
            logger.warning(f"Markdown 标题切片失败，回退到普通切片: {e}")
            return None  # 回退到普通切片

        # 如果某个 MD chunk 仍然太长，再用 RecursiveCharacterTextSplitter 二次切割
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "；", "，", " ", ""],
        )
        for md_chunk in md_chunks:
            if len(md_chunk.page_content) > chunk_size:
                sub_chunks = text_splitter.split_documents([md_chunk])
                # 继承标题 metadata
                for sub in sub_chunks:
                    sub.metadata.update(md_chunk.metadata)
                all_chunks.extend(sub_chunks)
            else:
                all_chunks.append(md_chunk)

    return all_chunks


def split_documents(docs: list, chunk_size: int = 800, chunk_overlap: int = 200, filename: str = None) -> list:
    """
    文档分块（V3 智能分块策略）

    优化点：
    - Markdown 文件优先使用 MarkdownHeaderTextSplitter（按标题层级切分，保留标题结构）
    - 非 Markdown 文件使用 RecursiveCharacterTextSplitter
    - chunk_size 800：保证一个完整的语义单元不被切断
    - chunk_overlap 200：25% 重叠率，确保跨 chunk 的信息不会丢失
    - 碎片合并：相邻短 chunk 自动合并，提升语义完整性
    - 结构标记：每个 chunk 标记类型（heading/table/list/paragraph）
    """
    # ===== Markdown 文件：使用标题层级切片 =====
    is_markdown = filename and filename.lower().endswith('.md')

    if is_markdown:
        md_chunks = _split_markdown_by_headers(docs, chunk_size, chunk_overlap)
        if md_chunks is not None:
            # MD 切片成功，给每个 chunk 分配索引
            for i, chunk in enumerate(md_chunks):
                chunk.metadata["chunk_index"] = i
                chunk.metadata["chunk_type"] = "md_section"
                # 把标题层级信息拼成可读的章节路径，方便检索时展示
                header_path = " > ".join(
                    v for k, v in sorted(chunk.metadata.items()) if k.startswith("Header")
                )
                if header_path:
                    chunk.metadata["section_path"] = header_path
            logger.info(f"Markdown 切片完成: {filename}, 共 {len(md_chunks)} 个标题级 chunk")
            return md_chunks
        # MD 切片失败，回退到普通切片
        logger.info(f"Markdown 标题切片回退到普通切片: {filename}")

    # ===== 非 Markdown 文件：使用 RecursiveCharacterTextSplitter =====
    separators = [
        "\n# ",      # Markdown 一级标题
        "\n## ",     # Markdown 二级标题
        "\n### ",    # Markdown 三级标题
        "\n\n",      # 段落分隔
        "\n",        # 换行
        "。",        # 中文句号
        "；",        # 中文分号
        "，",        # 中文逗号
        " ",         # 空格
        "",          # 最后按字符切
    ]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
    )
    chunks = splitter.split_documents(docs)

    # 后处理：合并碎片 chunk
    merged_chunks = []
    i = 0
    while i < len(chunks):
        current = chunks[i]
        current_text = current.page_content.strip()

        # 如果当前 chunk 太短（<200字），尝试与下一个合并
        if len(current_text) < 200 and i + 1 < len(chunks):
            next_chunk = chunks[i + 1]
            combined = current_text + "\n\n" + next_chunk.page_content.strip()
            if len(combined) <= chunk_size * 1.2:
                current.page_content = combined
                current.metadata["merged_with_next"] = True
                i += 2
                merged_chunks.append(current)
                continue

        # 标记 chunk 的结构类型
        current_text = current.page_content.strip()
        if current_text.startswith('#'):
            current.metadata["chunk_type"] = "heading"
        elif '|' in current_text and current_text.count('|') >= 3:
            current.metadata["chunk_type"] = "table"
        elif any(current_text.startswith(f'{n}.') or current_text.startswith(f'{n}、') for n in range(1, 10)):
            current.metadata["chunk_type"] = "list"
        else:
            current.metadata["chunk_type"] = "paragraph"

        merged_chunks.append(current)
        i += 1

    # 给每个 chunk 分配唯一 ID
    for i, chunk in enumerate(merged_chunks):
        chunk.metadata["chunk_index"] = i

    return merged_chunks


def index_document(file_path: str, filename: str = None, agent_id: str = None) -> dict:
    """
    完整的文档索引流程：加载 → 分块 → 索引存储

    [#11] 智能降级策略：
    - 当 Embedding 可用时：使用 ChromaDB 向量索引（语义搜索 + 关键词搜索）
    - 当 Embedding 不可用时：自动降级为关键词索引（纯关键词搜索）
    - 降级过程完全自动，对上层调用者透明

    Returns:
        dict: 包含分块数量和状态信息
    """
    global _embedding_available

    if filename is None:
        filename = os.path.basename(file_path)
    # URL解码文件名：LLM传入的文件名可能是URL编码的，统一解码为中文
    filename = unquote(filename)

    logger.info(f"开始索引文档: {filename}")

    # 1. 加载文档
    docs = load_document(file_path)

    # 2. 给文档添加元数据
    for doc in docs:
        doc.metadata["source_file"] = filename

    # 3. 分块（传入 filename，MD 文件自动使用标题层级切片）
    chunks = split_documents(docs, filename=filename)

    if not chunks:
        return {
            "filename": filename,
            "chunks": 0,
            "status": "success",
            "message": f"文档 {filename} 内容为空，无需索引",
        }

    total_chunks = len(chunks)

    # ===== [#11] 智能索引：向量优先，关键词降级 =====

    # 情况1：已知 Embedding 不可用 → 直接使用关键词索引
    if _embedding_available is False:
        _add_chunks_to_keyword_index(chunks, filename, agent_id)
        _bm25_cache_invalidation(agent_id)
        return {
            "filename": filename,
            "chunks": total_chunks,
            "status": "success",
            "indexing_mode": "keyword",
            "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- Embedding API 不可用，已自动切换为关键词搜索",
        }

    # 情况2：尝试向量索引
    try:
        vector_store = get_vector_store(agent_id=agent_id)

        if vector_store is None:
            # Embedding 不可用，降级为关键词索引
            _add_chunks_to_keyword_index(chunks, filename, agent_id)
            return {
                "filename": filename,
                "chunks": total_chunks,
                "status": "success",
                "indexing_mode": "keyword",
                "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- Embedding 不可用，已自动切换为关键词搜索",
            }

        # 尝试分批向量化并存储
        batch_count = (total_chunks + EMBEDDING_BATCH_SIZE - 1) // EMBEDDING_BATCH_SIZE

        for i in range(batch_count):
            start = i * EMBEDDING_BATCH_SIZE
            end = min(start + EMBEDDING_BATCH_SIZE, total_chunks)
            batch = chunks[start:end]
            try:
                vector_store.add_documents(batch)
            except Exception as e:
                error_str = str(e)

                if _is_embedding_error(e):
                    # [#11] Embedding 不可用 → 自动降级为关键词索引
                    logger.warning(f"Embedding 不可用（{error_str}），自动切换为关键词索引模式")

                    # 标记为不可用
                    _embedding_available = False
                    _embedding_degraded_at = time.time()  # [优化5] 记录降级时间

                    # 尝试回滚已写入该文档的数据
                    try:
                        collection = vector_store._collection
                        existing = collection.get(
                            where={"source_file": filename},
                            include=["metadatas"],
                        )
                        if existing.get("ids"):
                            collection.delete(ids=existing["ids"])
                    except Exception:
                        pass

                    # 降级：将全部分块（不仅是当前批次）写入关键词索引
                    _add_chunks_to_keyword_index(chunks, filename, agent_id)

                    return {
                        "filename": filename,
                        "chunks": total_chunks,
                        "status": "success",
                        "indexing_mode": "keyword",
                        "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- Embedding API 不可用，已自动切换为关键词搜索",
                    }
                else:
                    # 非 Embedding 错误（如 ChromaDB 内部错误），仍然尝试降级
                    logger.error(f"向量化失败（非Embedding错误）: {error_str}")

                    # 回滚已写入的数据
                    try:
                        collection = vector_store._collection
                        existing = collection.get(
                            where={"source_file": filename},
                            include=["metadatas"],
                        )
                        if existing.get("ids"):
                            collection.delete(ids=existing["ids"])
                    except Exception:
                        pass

                    # 降级为关键词索引
                    _embedding_available = False
                    _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
                    _add_chunks_to_keyword_index(chunks, filename, agent_id)

                    return {
                        "filename": filename,
                        "chunks": total_chunks,
                        "status": "success",
                        "indexing_mode": "keyword",
                        "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- 向量化失败，已自动切换为关键词搜索",
                    }

        # 向量索引成功
        logger.info(f"文档索引完成（向量模式）: {filename}, 共 {total_chunks} 个分块")
        _embedding_available = True
        _bm25_cache_invalidation(agent_id)  # 清除BM25缓存，确保新文档可被搜索

        # 清除BM25缓存（新文档索引后，旧缓存不再完整）
        cache_key = agent_id or "__global__"
        if cache_key in _bm25_doc_cache:
            del _bm25_doc_cache[cache_key]

        return {
            "filename": filename,
            "chunks": total_chunks,
            "status": "success",
            "indexing_mode": "vector",
            "message": f"文档 {filename} 已成功索引（向量模式，共 {total_chunks} 个分块，分 {batch_count} 批写入）",
        }

    except Exception as e:
        # 整个向量流程异常，降级为关键词索引
        logger.error(f"向量索引流程异常: {e}")
        _embedding_available = False
        _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
        _add_chunks_to_keyword_index(chunks, filename, agent_id)

        return {
            "filename": filename,
            "chunks": total_chunks,
            "status": "success",
            "indexing_mode": "keyword",
            "message": f"文档 {filename} 已成功索引（关键词模式，共 {total_chunks} 个分块）- 向量索引异常，已自动切换为关键词搜索",
        }


# ===== [#9] 混合检索 + 重排序 =====
# ===== [#12] BM25 索引优化：rank_bm25 替代全量遍历 =====

# BM25 索引缓存：rank_bm25.BM25Okapi 实例 + 文档元数据
# 结构: { cache_key: { "bm25": BM25Okapi, "corpus": [str], "metadatas": [dict], "ids": [str], "updated_at": float } }
_bm25_index_cache = {}
_BM25_INDEX_TTL = 1800  # [优化7] 索引缓存30分钟（原5分钟，延长后减少空闲后索引重建 200ms-1s）
_BM25_INDEX_MAX_ENTRIES = 50  # [性能修复] 最多缓存50个智能体的BM25索引，超过则淘汰最久未使用的

# 旧版全量文档缓存（保留作为 rank_bm25 不可用时的降级方案）
_bm25_doc_cache = {}  # agent_id -> {"data": all_docs, "updated_at": float}
_BM25_CACHE_TTL = 120
_BM25_DOC_MAX_ENTRIES = 50  # [性能修复] 最多缓存50个智能体的文档缓存

# [性能优化 1] 关键词索引 BM25 缓存（_search_keyword_index 使用）
# 与 _bm25_index_cache 分离：_bm25_index_cache 从 ChromaDB 构建，
# 而 _keyword_bm25_cache 从 JSON 关键词索引文件构建，服务于降级/纯关键词模式
# 缓存结构: { cache_key: {"bm25": BM25Okapi, "corpus": [...], "index_data": [...], "updated_at": float, "index_hash": str} }
_keyword_bm25_cache = {}
_KEYWORD_BM25_TTL = 1800  # 30分钟，与 _BM25_INDEX_TTL 一致
_KEYWORD_BM25_MAX_ENTRIES = 50


def _tokenize_text(text: str) -> list[str]:
    """[#12] 中文分词：优先 jieba，回退到正则切词

    jieba 对中文的分词质量远高于简单的正则匹配，能正确处理如：
    - "动力电池系统" → ["动力", "电池", "系统"]
    - "故障诊断方法" → ["故障", "诊断", "方法"]
    而正则只能整句匹配 "动力电池系统"，导致 BM25 无法匹配部分关键词。
    """
    if _JIEBA_AVAILABLE:
        # jieba 分词 + 去停用词 + 转小写
        tokens = jieba.lcut(text.lower())
        return [t for t in tokens if t.strip() and t not in _STOPWORDS]
    else:
        # 回退：正则提取中文词和英文词
        tokens = re.findall(r'[\u4e00-\u9fff]+|\w+', text.lower())
        return [t for t in tokens if t not in _STOPWORDS]


def _build_bm25_index(agent_id: str = None) -> dict:
    """[#12] 构建 rank_bm25 索引（带缓存）

    从 ChromaDB 获取全量文档，构建 BM25Okapi 倒排索引。
    索引构建后缓存在内存中，后续搜索直接走倒排索引，无需全量遍历。

    优势：
    - 旧版：每次搜索遍历所有文档计算 TF 分数，O(N) 复杂度
    - 新版：构建一次倒排索引，后续搜索走倒排链，O(K) 复杂度（K为命中文档数）

    Returns:
        dict: {"bm25": BM25Okapi, "corpus": [...], "metadatas": [...], "ids": [...]}
              如果构建失败返回 None
    """
    cache_key = agent_id or "__global__"
    now = time.time()

    # 检查缓存是否有效
    if cache_key in _bm25_index_cache:
        entry = _bm25_index_cache[cache_key]
        if now - entry["updated_at"] < _BM25_INDEX_TTL:
            return entry

    # 从 ChromaDB 获取全量文档
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is None:
        return None

    try:
        collection = vector_store._collection
        all_docs = collection.get(include=["documents", "metadatas"])
    except Exception as e:
        logger.warning(f"获取全量文档失败: {e}")
        return None

    if not all_docs.get("ids"):
        return None

    # 分词构建语料库
    corpus = []
    metadatas = []
    ids = []
    for i, doc_id in enumerate(all_docs["ids"]):
        content = all_docs["documents"][i] or ""
        if not content.strip():
            continue
        corpus.append(content)
        metadatas.append(all_docs["metadatas"][i] or {})
        ids.append(doc_id)

    if not corpus:
        return None

    # 分词
    tokenized_corpus = [_tokenize_text(doc) for doc in corpus]

    # 构建 BM25Okapi 索引
    try:
        bm25 = BM25Okapi(tokenized_corpus)
    except Exception as e:
        logger.error(f"BM25 索引构建失败: {e}")
        return None

    index_data = {
        "bm25": bm25,
        "corpus": corpus,
        "metadatas": metadatas,
        "ids": ids,
        "updated_at": now,
    }
    _bm25_index_cache[cache_key] = index_data
    # [性能修复] LRU 淘汰：超过最大条目数时，移除最久未更新的条目
    while len(_bm25_index_cache) > _BM25_INDEX_MAX_ENTRIES:
        oldest_key = min(_bm25_index_cache, key=lambda k: _bm25_index_cache[k]["updated_at"])
        del _bm25_index_cache[oldest_key]
        logger.info(f"[#12] BM25 索引缓存 LRU 淘汰: {oldest_key}")
    logger.info(f"[#12] BM25 索引已构建: cache_key={cache_key}, 文档数={len(corpus)}")
    return index_data


def _get_all_docs_cached(agent_id: str = None) -> dict:
    """从 ChromaDB 获取全量文档（带缓存，避免重复IO）

    缓存策略：每个 agent_id 的全量文档数据缓存 _BM25_CACHE_TTL 秒，
    避免每次搜索都执行 collection.get() 这个重量级操作。

    [#12] 此函数仅作为 rank_bm25 不可用时的降级方案，主流程已改用 _build_bm25_index()
    """
    cache_key = agent_id or "__global__"
    now = time.time()

    if cache_key in _bm25_doc_cache:
        entry = _bm25_doc_cache[cache_key]
        if now - entry["updated_at"] < _BM25_CACHE_TTL:
            return entry["data"]

    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is None:
        return {}

    try:
        collection = vector_store._collection
        all_docs = collection.get(include=["documents", "metadatas"])
        _bm25_doc_cache[cache_key] = {"data": all_docs, "updated_at": now}
        # [性能修复] LRU 淘汰：超过最大条目数时，移除最久未更新的条目
        while len(_bm25_doc_cache) > _BM25_DOC_MAX_ENTRIES:
            oldest_key = min(_bm25_doc_cache, key=lambda k: _bm25_doc_cache[k]["updated_at"])
            del _bm25_doc_cache[oldest_key]
        return all_docs
    except Exception as e:
        logger.warning(f"获取全量文档失败: {e}")
        return {}


def _bm25_cache_invalidation(agent_id: str = None):
    """清除BM25缓存（文档变更时调用）

    [#12] 同时清除 rank_bm25 索引缓存和旧版全量文档缓存
    """
    global _bm25_index_cache, _bm25_doc_cache
    cache_key = agent_id or "__global__"
    cleared = []
    if cache_key in _bm25_index_cache:
        del _bm25_index_cache[cache_key]
        cleared.append("rank_bm25索引")
    if cache_key in _bm25_doc_cache:
        del _bm25_doc_cache[cache_key]
        cleared.append("全量文档缓存")
    if cleared:
        logger.info(f"BM25缓存已清除: {cache_key} ({', '.join(cleared)})")


def cleanup_bm25_caches():
    """[性能修复] 定期清理过期/过多的 BM25 缓存条目，由 main.py 的定期任务调用
    
    清理策略：
    1. 淘汰超过 TTL 的 BM25 索引缓存条目
    2. 淘汰超过 TTL 的全量文档缓存条目
    3. 如果缓存总数超过最大条目数，淘汰最老的
    """
    now = time.time()
    
    # 清理过期的 BM25 索引缓存
    stale_index = [k for k, v in _bm25_index_cache.items()
                   if now - v.get("updated_at", 0) > _BM25_INDEX_TTL]
    for k in stale_index:
        del _bm25_index_cache[k]
    
    # 清理过期的全量文档缓存
    stale_doc = [k for k, v in _bm25_doc_cache.items()
                 if now - v.get("updated_at", 0) > _BM25_CACHE_TTL]
    for k in stale_doc:
        del _bm25_doc_cache[k]
    
    # 如果仍然超过最大条目数，淘汰最老的
    while len(_bm25_index_cache) > _BM25_INDEX_MAX_ENTRIES:
        oldest = min(_bm25_index_cache, key=lambda k: _bm25_index_cache[k].get("updated_at", 0))
        del _bm25_index_cache[oldest]
        stale_index.append(oldest)
    
    while len(_bm25_doc_cache) > _BM25_DOC_MAX_ENTRIES:
        oldest = min(_bm25_doc_cache, key=lambda k: _bm25_doc_cache[k].get("updated_at", 0))
        del _bm25_doc_cache[oldest]
        stale_doc.append(oldest)
    
    # [性能优化 1] 清理过期的关键词 BM25 缓存
    stale_kw = [k for k, v in _keyword_bm25_cache.items()
                if now - v.get("updated_at", 0) > _KEYWORD_BM25_TTL]
    for k in stale_kw:
        del _keyword_bm25_cache[k]
    
    while len(_keyword_bm25_cache) > _KEYWORD_BM25_MAX_ENTRIES:
        oldest = min(_keyword_bm25_cache, key=lambda k: _keyword_bm25_cache[k].get("updated_at", 0))
        del _keyword_bm25_cache[oldest]
        stale_kw.append(oldest)
    
    total = len(stale_index) + len(stale_doc) + len(stale_kw)
    if total > 0:
        logger.info(f"[定期清理] BM25缓存清理: 索引{len(stale_index)}条, 文档{len(stale_doc)}条, 关键词{len(stale_kw)}条, 剩余索引={len(_bm25_index_cache)}, 文档={len(_bm25_doc_cache)}, 关键词={len(_keyword_bm25_cache)}")


def _bm25_keyword_search(query: str, top_k: int = 10, agent_id: str = None) -> list[dict]:
    """
    [#9] BM25 关键词检索
    [#12] 优化：使用 rank_bm25.BM25Okapi 倒排索引替代全量遍历

    性能提升：
    - 旧版：O(N) 全量遍历，每次搜索都要遍历所有文档
    - 新版：O(K) 倒排索引搜索，只遍历命中文档（K << N）
    - 中文分词：jieba 分词替代简单正则，显著提升中文召回率

    当 rank_bm25 未安装时自动降级为旧版全量遍历模式
    当 ChromaDB 不可用时返回空列表（由 _search_keyword_index 替代）
    """
    # ===== [#12] 优先使用 rank_bm25 索引 =====
    if _RANK_BM25_AVAILABLE:
        try:
            index_data = _build_bm25_index(agent_id=agent_id)
            if index_data is not None:
                bm25 = index_data["bm25"]
                corpus = index_data["corpus"]
                metadatas = index_data["metadatas"]
                ids = index_data["ids"]

                # 对查询进行分词
                tokenized_query = _tokenize_text(query)
                if not tokenized_query:
                    return []

                # 使用 BM25Okapi 的 get_scores 方法获取所有文档分数
                # 内部走倒排索引，只计算包含查询词的文档
                doc_scores = bm25.get_scores(tokenized_query)

                # 收集得分 > 0 的结果
                scored = []
                for i, score in enumerate(doc_scores):
                    if score <= 0:
                        continue
                    scored.append({
                        "content": corpus[i],
                        "source": metadatas[i].get("source_file", "未知来源"),
                        "chunk_index": metadatas[i].get("chunk_index", -1),
                        "bm25_score": round(float(score), 4),
                        "id": ids[i],
                    })

                # 按 BM25 分排序
                scored.sort(key=lambda x: x["bm25_score"], reverse=True)
                return scored[:top_k]
        except Exception as e:
            logger.warning(f"[#12] rank_bm25 搜索失败，降级为全量遍历: {e}")

    # ===== 降级：旧版全量遍历（rank_bm25 未安装或出错时） =====
    all_docs = _get_all_docs_cached(agent_id=agent_id)

    if not all_docs.get("ids"):
        return []

    query_terms = set(re.findall(r'[\u4e00-\u9fff]+|\w+', query.lower()))
    query_terms = query_terms - _STOPWORDS

    scored = []
    for i, doc_id in enumerate(all_docs["ids"]):
        content = all_docs["documents"][i] or ""
        metadata = all_docs["metadatas"][i] or {}

        # 计算关键词匹配分
        content_lower = content.lower()
        match_count = sum(1 for term in query_terms if term in content_lower)
        if match_count == 0:
            continue

        # TF 近似：关键词出现次数 / 文档长度
        tf_score = match_count / max(len(content), 1) * 1000

        scored.append({
            "content": content,
            "source": metadata.get("source_file", "未知来源"),
            "chunk_index": metadata.get("chunk_index", -1),
            "bm25_score": tf_score,
            "id": doc_id,
        })

    # 按 BM25 分排序
    scored.sort(key=lambda x: x["bm25_score"], reverse=True)
    return scored[:top_k]


def _reciprocal_rank_fusion(vector_results: list[dict], keyword_results: list[dict], k: int = 60) -> list[dict]:
    """
    [#9] 倒数排名融合（Reciprocal Rank Fusion）
    将向量检索和关键词检索的结果融合，按融合分数排序

    RRF公式: score = 1/(k + rank_vector) + 1/(k + rank_keyword)
    """
    fused_scores = {}

    # 向量检索结果
    for rank, item in enumerate(vector_results):
        content_key = item["content"][:200]  # 用内容前200字作为唯一标识
        if content_key not in fused_scores:
            fused_scores[content_key] = {**item, "rrf_score": 0}
        fused_scores[content_key]["rrf_score"] += 1.0 / (k + rank + 1)
        # 保留向量相似度
        if "relevance_score" not in fused_scores[content_key]:
            fused_scores[content_key]["relevance_score"] = item.get("relevance_score", 0)

    # 关键词检索结果
    for rank, item in enumerate(keyword_results):
        content_key = item["content"][:200]
        if content_key not in fused_scores:
            fused_scores[content_key] = {
                "content": item["content"],
                "source": item.get("source", "未知来源"),
                "chunk_index": item.get("chunk_index", -1),
                "relevance_score": 0,
                "rrf_score": 0,
            }
        fused_scores[content_key]["rrf_score"] += 1.0 / (k + rank + 1)
        # 如果有 BM25 分，补充
        if "bm25_score" in item and "bm25_score" not in fused_scores[content_key]:
            fused_scores[content_key]["bm25_score"] = item["bm25_score"]

    # 按 RRF 分排序
    results = sorted(fused_scores.values(), key=lambda x: x["rrf_score"], reverse=True)
    return results


def _generate_multi_queries(query: str) -> list[str]:
    """为查询生成多个变体，提升短段落检索召回率
    
    当用户查询某部门职责时，原文中可能用不同表述（如"设备部"vs"8 设备部"），
    多 query 检索可以覆盖不同表述，避免短职责段落被长过程描述淹没。
    """
    queries = [query]
    
    # 部门职责类查询：扩展搜索词
    dept_patterns = [
        (r'(\w+)部(?:的)?职[责责任务]', lambda m: [m.group(0), m.group(1) + '部']),
        (r'(\w+)部门(?:的)?职[责责任务]', lambda m: [m.group(0), m.group(1) + '部门']),
    ]
    import re
    for pattern, gen in dept_patterns:
        match = re.search(pattern, query)
        if match:
            queries.extend(gen(match))
            break
    
    return list(dict.fromkeys(queries))  # 去重保序

def search_documents(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """
    [#9] 混合检索：向量语义检索 + BM25关键词检索 + RRF融合
    [#10] 引用溯源：返回结果标注文档名 + 段落位置
    [#11] 自动降级：Embedding 不可用时仅使用关键词检索
    [#13] 多查询检索：对部门职责等查询自动扩展搜索词，提升短段落召回率

    Args:
        query: 用户查询
        top_k: 返回最相关的 K 个结果
        agent_id: 智能体ID

    Returns:
        list[dict]: 检索结果列表
    """
    global _embedding_available

    # ===== 普通聊天模式（无 agent_id）：无知识库，返回空结果 =====
    if not agent_id:
        logger.info(f"普通聊天模式无知识库，跳过检索: query='{query[:50]}...'")
        return []

    # ===== [#11] 根据索引模式选择检索策略 =====

    if _embedding_available is False:
        # 关键词模式：仅使用关键词索引检索
        logger.info(f"关键词模式检索: query='{query[:50]}...', agent_id={agent_id}")
        results = _search_keyword_index(query, top_k=top_k, agent_id=agent_id)
        if not results:
            # 关键词索引无结果，尝试从磁盘文件全文搜索
            results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)
        return results

    # ===== [#13] 多查询检索：生成查询变体，提升短段落召回率 =====
    queries = _generate_multi_queries(query)
    if len(queries) > 1:
        logger.info(f"多查询检索: {queries}")

    # ===== 向量模式：混合检索 =====

    # 1. 向量语义检索 - 对每个查询变体都检索，合并去重
    vector_store = get_vector_store(agent_id=agent_id)
    vector_results_raw = []
    vector_results = []

    if vector_store is not None:
        seen_contents = set()
        for q in queries:
            try:
                raw = vector_store.similarity_search_with_score(q, k=top_k * 2)
                for doc, score in raw:
                    # 用内容前100字去重
                    content_key = doc.page_content[:100]
                    if content_key not in seen_contents:
                        seen_contents.add(content_key)
                        vector_results_raw.append((doc, score))
            except Exception as e:
                error_str = str(e)
                logger.warning(f"向量检索失败: {e}")
                if _is_embedding_error(e):
                    logger.warning(f"Embedding API 不可用，自动切换为关键词检索模式")
                    _embedding_available = False
                    _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
                    results = _search_keyword_index(query, top_k=top_k, agent_id=agent_id)
                    if not results:
                        results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)
                    return results
        
        # 按相似度排序，取 top_k * 3
        vector_results_raw.sort(key=lambda x: x[1])
        vector_results_raw = vector_results_raw[:top_k * 3]
        
        for doc, score in vector_results_raw:
            vector_results.append({
                "content": doc.page_content,
                "source": doc.metadata.get("source_file", "未知来源"),
                "chunk_index": doc.metadata.get("chunk_index", -1),
                "relevance_score": round(1 - score, 4),
            })
    else:
        # vector_store 为 None，降级为关键词模式
        _embedding_available = False
        _embedding_degraded_at = time.time()  # [优化5] 记录降级时间
        results = _search_keyword_index(query, top_k=top_k, agent_id=agent_id)
        if not results:
            results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)
        return results

    # 2. BM25 关键词检索 - 对每个查询变体都检索，合并去重
    keyword_results = []
    try:
        seen_kw = set()
        for q in queries:
            kw_raw = _bm25_keyword_search(q, top_k=top_k * 2, agent_id=agent_id)
            for item in kw_raw:
                content_key = item.get("content", "")[:100]
                if content_key not in seen_kw:
                    seen_kw.add(content_key)
                    keyword_results.append(item)
        keyword_results = keyword_results[:top_k * 3]
    except Exception as e:
        logger.warning(f"BM25关键词检索失败，跳过: {e}")

    # 3. RRF 融合
    if keyword_results:
        fused_results = _reciprocal_rank_fusion(vector_results, keyword_results)
    else:
        # 关键词检索失败，直接用向量结果
        fused_results = vector_results

    # 4. 如果融合结果为空，尝试磁盘文件兜底搜索
    if not fused_results:
        logger.info(f"向量+关键词均无结果，尝试磁盘文件搜索: query='{query[:50]}...'")
        fused_results = _search_disk_files(query, top_k=top_k, agent_id=agent_id)

    # 5. 上下文窗口增强：检索到 chunk 后，自动补全其前后相邻 chunk
    # 这样 AI 可以看到完整的前后文，避免"断章取义"
    formatted = _expand_context_window(fused_results[:top_k], agent_id=agent_id)

    return formatted


def _expand_context_window(results: list[dict], agent_id: str = None, window_size: int = 1) -> list[dict]:
    """上下文窗口增强：为检索结果补全前后相邻 chunk

    当检索到某个 chunk 时，它的前后 chunk 可能包含关键上下文。
    例如检索到"步骤3"，但步骤1-2在另一个chunk中，AI不知道前提条件。

    策略：
    - 对每个检索结果，查找其 chunk_index ± window_size 的相邻 chunk
    - 将相邻 chunk 的内容拼接到当前结果中（标注为上下文）
    - 避免重复：如果两个检索结果的上下文重叠，只保留一次

    Args:
        results: 原始检索结果列表
        agent_id: 智能体ID
        window_size: 前后扩展几个 chunk（默认1，即前1后1）

    Returns:
        list[dict]: 增强后的检索结果
    """
    if not results:
        return results

    # 收集所有需要查询的 (source_file, chunk_index) 对
    expanded_results = []
    seen_context_keys = set()  # 避免重复上下文

    for r in results:
        source = r.get("source", "")
        chunk_idx = r.get("chunk_index", -1)
        content = r.get("content", "")

        if not source or chunk_idx < 0:
            # 无法扩展上下文，直接返回
            expanded_results.append(r)
            continue

        # 查找前后 chunk 的内容
        context_before = ""
        context_after = ""

        # 从向量数据库获取同一文档的相邻 chunk
        vector_store = get_vector_store(agent_id=agent_id)
        if vector_store is not None:
            try:
                collection = vector_store._collection
                # 查找同一文档中 chunk_index 在 [chunk_idx-window, chunk_idx+window] 范围内的 chunk
                nearby_chunks = collection.get(
                    where={
                        "$and": [
                            {"source_file": source},
                            {"chunk_index": {"$gte": chunk_idx - window_size}},
                            {"chunk_index": {"$lte": chunk_idx + window_size}},
                        ]
                    },
                    include=["documents", "metadatas"],
                )

                if nearby_chunks and nearby_chunks.get("ids"):
                    for i, meta in enumerate(nearby_chunks["metadatas"]):
                        nearby_idx = meta.get("chunk_index", -1)
                        nearby_content = nearby_chunks["documents"][i] or ""

                        if nearby_idx == chunk_idx:
                            continue  # 跳过自身
                        context_key = f"{source}:{nearby_idx}"
                        if context_key in seen_context_keys:
                            continue

                        if nearby_idx < chunk_idx:
                            context_before += nearby_content + "\n"
                            seen_context_keys.add(context_key)
                        elif nearby_idx > chunk_idx:
                            context_after += "\n" + nearby_content
                            seen_context_keys.add(context_key)
            except Exception as e:
                logger.debug(f"上下文窗口扩展失败（不影响主流程）: {e}")

        # 组装增强后的内容
        enhanced_content = ""
        if context_before.strip():
            enhanced_content += f"[上文参考] {context_before.strip()}\n\n"
        enhanced_content += content
        if context_after.strip():
            enhanced_content += f"\n\n[下文参考] {context_after.strip()}"

        expanded_results.append({
            "content": enhanced_content,
            "source": r.get("source", "未知来源"),
            "chunk_index": r.get("chunk_index", -1),
            "relevance_score": r.get("relevance_score", 0),
        })

    return expanded_results


def _search_disk_files(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """[#11] 磁盘文件全文搜索（关键词索引的补充）

    当关键词索引中也没有匹配结果时，直接读取磁盘上的文档文件做搜索。
    这是最后的兜底方案，确保即使没有任何索引，用户也能查到文档内容。

    Args:
        query: 查询文本
        top_k: 返回结果数
        agent_id: 智能体ID

    Returns:
        list[dict]: 搜索结果
    """
    if agent_id:
        scan_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")
    else:
        scan_dir = settings.DOCUMENTS_DIR

    if not os.path.exists(scan_dir):
        return []

    query_terms = set(re.findall(r'[\u4e00-\u9fff]+|\w+', query.lower()))
    query_terms = query_terms - _STOPWORDS

    if not query_terms:
        return []

    scored = []
    for fname in os.listdir(scan_dir):
        ext = os.path.splitext(fname)[1].lower()
        if ext not in {'.txt', '.docx', '.pdf'}:
            continue
        file_path = os.path.join(scan_dir, fname)
        if not os.path.isfile(file_path):
            continue

        try:
            docs = load_document(file_path)
            for doc in docs:
                content = doc.page_content
                content_lower = content.lower()
                match_count = sum(1 for term in query_terms if term in content_lower)
                if match_count > 0:
                    term_coverage = match_count / max(len(query_terms), 1)
                    tf_score = match_count / max(len(content), 1) * 1000
                    combined_score = tf_score * 0.6 + term_coverage * 100 * 0.4
                    scored.append({
                        "content": content[:2000],  # 限制长度避免过大
                        "source": fname,
                        "chunk_index": 0,
                        "relevance_score": round(combined_score, 4),
                    })
        except Exception:
            continue

    scored.sort(key=lambda x: x["relevance_score"], reverse=True)
    return scored[:top_k]


def get_document_content(filename: str, agent_id: str = None) -> dict:
    """获取知识库中指定文档的完整内容（从磁盘原始文件读取，不依赖向量搜索）

    与 search_documents 不同，此函数返回文档的完整文本内容，
    而不是分块后的片段。用于文档修改前获取完整内容。

    Args:
        filename: 文档文件名（含扩展名）
        agent_id: 智能体ID（用于验证文档归属，不参与向量搜索）

    Returns:
        dict: 包含文档完整内容、状态信息
    """
    # 查找文件路径（可能在agent子目录中）
    if agent_id:
        file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)
        if not os.path.exists(file_path):
            file_path = os.path.join(settings.DOCUMENTS_DIR, filename)
    else:
        file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

    if not os.path.exists(file_path):
        return {
            "filename": filename,
            "status": "not_found",
            "content": "",
            "message": f"文档 {filename} 在服务器上未找到",
        }

    try:
        docs = load_document(file_path)
        full_content = "\n".join([doc.page_content for doc in docs])

        if not full_content.strip():
            return {
                "filename": filename,
                "status": "empty",
                "content": "",
                "message": f"文档 {filename} 内容为空",
            }

        return {
            "filename": filename,
            "status": "success",
            "content": full_content,
            "char_count": len(full_content),
            "message": f"成功获取文档 {filename} 的完整内容（共 {len(full_content)} 字符）",
        }
    except Exception as e:
        return {
            "filename": filename,
            "status": "error",
            "content": "",
            "message": f"读取文档失败: {str(e)}",
        }


def list_indexed_documents(agent_id: str = None) -> list[str]:
    """列出知识库中所有已索引的文档（按 agent_id 隔离）

    [#11] 同时检查向量索引、关键词索引和磁盘文件，合并结果
    
    三重数据源合并策略：
    1. ChromaDB 向量索引 - 已成功向量化的文档
    2. 关键词索引 - 因 Embedding 不可用而降级为关键词索引的文档
    3. 磁盘文件扫描 - 兜底：确保即使索引丢失，文件仍然可见
    
    这保证了前端知识库面板、智能体搜索工具、API 列表接口三者看到一致的文档列表。
    
    注意：普通聊天模式（agent_id=None）没有知识库，返回空列表
    """
    # 普通聊天模式：无知识库
    if not agent_id:
        return []

    sources = set()

    # 1. 从 ChromaDB 获取（如果可用）
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            all_docs = collection.get(include=["metadatas"])
            for meta in all_docs["metadatas"]:
                if meta and "source_file" in meta:
                    sources.add(meta["source_file"])
        except Exception:
            pass

    # 2. 从关键词索引获取
    keyword_docs = _load_keyword_index(agent_id)
    for entry in keyword_docs:
        if entry.get("source_file"):
            sources.add(entry["source_file"])

    # 3. 磁盘文件扫描（兜底：确保文件存在但索引丢失时仍可见）
    if agent_id:
        scan_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")
    else:
        scan_dir = settings.DOCUMENTS_DIR
    if os.path.exists(scan_dir):
        for fname in os.listdir(scan_dir):
            ext = os.path.splitext(fname)[1].lower()
            if ext in {'.pdf', '.txt', '.docx'}:
                file_path = os.path.join(scan_dir, fname)
                if os.path.isfile(file_path):
                    sources.add(fname)

    return sorted(list(sources))


def update_document(filename: str, new_content: str, agent_id: str = None, async_reindex: bool = False) -> dict:
    """
    修改知识库中已有文档的内容
    流程：删除旧的向量分块 → 用新内容覆盖原文件 → 重新索引

    [#11] 同时更新向量索引和关键词索引

    Args:
        filename: 要修改的文档文件名
        new_content: 新的文档内容（纯文本）
        agent_id: 智能体ID
        async_reindex: 是否异步重索引

    Returns:
        dict: 包含修改状态和详细信息
    """
    # 1. 检查文件是否存在（可能在agent子目录中）
    if agent_id:
        file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)
        if not os.path.exists(file_path):
            file_path = os.path.join(settings.DOCUMENTS_DIR, filename)
    else:
        file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

    if not os.path.exists(file_path):
        return {
            "filename": filename,
            "status": "not_found",
            "message": f"文档 {filename} 在服务器上未找到",
        }

    # 2. 删除旧的索引数据
    chunks_deleted = 0

    # 从 ChromaDB 删除
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            results = collection.get(
                where={"source_file": filename},
                include=["metadatas"],
            )
            chunk_ids = results.get("ids", [])
            if chunk_ids:
                collection.delete(ids=chunk_ids)
                chunks_deleted = len(chunk_ids)
        except Exception as e:
            logger.warning(f"删除旧向量分块时出错: {e}")

    # 从关键词索引删除
    keyword_deleted = _delete_from_keyword_index(filename, agent_id)
    _bm25_cache_invalidation(agent_id)  # 清除BM25缓存

    # 清除BM25缓存（文档内容已变更，旧缓存失效）
    cache_key = agent_id or "__global__"
    if cache_key in _bm25_doc_cache:
        del _bm25_doc_cache[cache_key]

    # 3. 用新内容覆盖原文件
    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".txt":
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(new_content)
        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                for line in new_content.split("\n"):
                    doc.add_paragraph(line)
                doc.save(file_path)
            except ImportError:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(new_content)
        elif ext == ".pdf":
            txt_path = file_path.rsplit('.', 1)[0] + '.txt'
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(new_content)
            os.remove(file_path)
            filename = filename.rsplit('.', 1)[0] + '.txt'
            file_path = txt_path
        else:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(new_content)
    except Exception as e:
        return {
            "filename": filename,
            "status": "error",
            "message": f"写入文件失败: {str(e)}",
        }

    # 4. 重新索引
    if async_reindex:
        import threading
        def _background_reindex(fp, fn, aid):
            try:
                index_result = index_document(fp, fn, agent_id=aid)
                logger.info(f"后台重索引完成: {fn}, {index_result.get('chunks', 0)} 个分块, 模式={index_result.get('indexing_mode', 'unknown')}")
            except Exception as e:
                logger.error(f"后台重索引失败: {fn}, {e}")

        thread = threading.Thread(target=_background_reindex, args=(file_path, filename, agent_id), daemon=True)
        thread.start()

        return {
            "filename": filename,
            "status": "success",
            "chunks_deleted": chunks_deleted,
            "keyword_entries_deleted": keyword_deleted,
            "chunks_indexed": "后台索引中",
            "message": f"文档 {filename} 已成功修改（删除 {chunks_deleted} 个向量分块 + {keyword_deleted} 个关键词条目，新内容正在后台索引中）",
        }
    else:
        try:
            index_result = index_document(file_path, filename, agent_id=agent_id)
        except Exception as e:
            return {
                "filename": filename,
                "status": "error",
                "message": f"重新索引失败: {str(e)}",
                "chunks_deleted": chunks_deleted,
            }

        mode = index_result.get("indexing_mode", "unknown")
        return {
            "filename": filename,
            "status": "success",
            "chunks_deleted": chunks_deleted,
            "keyword_entries_deleted": keyword_deleted,
            "chunks_indexed": index_result.get("chunks", 0),
            "indexing_mode": mode,
            "message": f"文档 {filename} 已成功修改（{mode}模式，删除 {chunks_deleted} 个向量分块 + {keyword_deleted} 个关键词条目，重新索引 {index_result.get('chunks', 0)} 个新分块）",
        }


def delete_document(filename: str, agent_id: str = None) -> dict:
    """
    从知识库中删除指定文档
    包括：从 ChromaDB 删除向量分块 + 从关键词索引删除 + 删除原始文件

    [#11] 同时清理向量索引和关键词索引

    Args:
        filename: 要删除的文档文件名
        agent_id: 智能体ID

    Returns:
        dict: 包含删除状态和详细信息
    """
    chunks_deleted = 0
    found_in_any = False

    # 1. 从 ChromaDB 删除
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            results = collection.get(
                where={"source_file": filename},
                include=["metadatas"],
            )
            chunk_ids = results.get("ids", [])
            if chunk_ids:
                found_in_any = True
                collection.delete(ids=chunk_ids)
                chunks_deleted = len(chunk_ids)
        except Exception as e:
            logger.warning(f"从 ChromaDB 删除失败: {e}")

    # 2. 从关键词索引删除
    keyword_deleted = _delete_from_keyword_index(filename, agent_id)
    if keyword_deleted > 0:
        found_in_any = True
    _bm25_cache_invalidation(agent_id)  # 清除BM25缓存

    # 清除BM25缓存（文档已删除，旧缓存失效）
    cache_key = agent_id or "__global__"
    if cache_key in _bm25_doc_cache:
        del _bm25_doc_cache[cache_key]

    # 3. 删除原始文件（查找可能的位置）
    file_deleted = False
    possible_paths = []

    if agent_id:
        possible_paths.append(os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename))
    possible_paths.append(os.path.join(settings.DOCUMENTS_DIR, filename))

    for file_path in possible_paths:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                file_deleted = True
                found_in_any = True
                break
            except Exception as e:
                return {
                    "filename": filename,
                    "chunks_deleted": chunks_deleted,
                    "keyword_entries_deleted": keyword_deleted,
                    "file_deleted": False,
                    "status": "partial",
                    "message": f"索引已删除，但原始文件删除失败: {str(e)}",
                }

    if not found_in_any:
        return {
            "filename": filename,
            "status": "not_found",
            "message": f"文档 {filename} 在知识库中未找到",
        }

    return {
        "filename": filename,
        "chunks_deleted": chunks_deleted,
        "keyword_entries_deleted": keyword_deleted,
        "file_deleted": file_deleted,
        "status": "success",
        "message": f"文档 {filename} 已成功删除（{chunks_deleted} 个向量分块 + {keyword_deleted} 个关键词条目，原始文件{'已删除' if file_deleted else '不存在'}）",
    }


def delete_agent_collection(agent_id: str) -> dict:
    """删除智能体的整个知识库 collection

    删除智能体时调用，清理以下内容：
    1. ChromaDB 中该智能体专属的 collection
    2. 内存中的向量数据库缓存
    3. 关键词索引文件
    4. 磁盘上的文档目录（data/documents/agent_{id}/）

    Args:
        agent_id: 智能体ID

    Returns:
        dict: 包含删除状态和详细信息
    """
    if not agent_id:
        return {"status": "error", "message": "agent_id 不能为空"}

    import shutil
    import chromadb
    collection_name = _get_collection_name(agent_id)
    cleanup_details = []

    try:
        # 1. 删除 ChromaDB collection
        client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
        existing_collections = [c.name for c in client.list_collections()]
        if collection_name in existing_collections:
            client.delete_collection(collection_name)
            logger.info(f"已删除智能体 ChromaDB collection: {collection_name}")
            cleanup_details.append("ChromaDB collection")

        # 2. 清理缓存
        cache_key = agent_id or "__global__"
        if cache_key in _vector_store_cache:
            del _vector_store_cache[cache_key]
            cleanup_details.append("向量缓存")

        # 3. 删除关键词索引文件
        keyword_index_path = _get_keyword_index_path(agent_id)
        if os.path.exists(keyword_index_path):
            try:
                os.remove(keyword_index_path)
                logger.info(f"已删除智能体关键词索引: {keyword_index_path}")
                cleanup_details.append("关键词索引")
            except Exception as e:
                logger.warning(f"删除关键词索引文件失败: {e}")

        # 4. 删除磁盘上的文档目录（解决孤立目录残留问题）
        agent_doc_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")
        if os.path.exists(agent_doc_dir):
            try:
                file_count = len([f for f in os.listdir(agent_doc_dir) if os.path.isfile(os.path.join(agent_doc_dir, f))])
                shutil.rmtree(agent_doc_dir)
                logger.info(f"已删除智能体文档目录: {agent_doc_dir} ({file_count} 个文件)")
                cleanup_details.append(f"文档目录({file_count}个文件)")
            except Exception as e:
                logger.warning(f"删除智能体文档目录失败: {e}")

        detail_str = " + ".join(cleanup_details) if cleanup_details else "无需清理"
        logger.info(f"智能体 {agent_id} 知识库清理完成: {detail_str}")
        return {"status": "success", "message": f"智能体知识库 {collection_name} 已删除（{detail_str}）"}
    except Exception as e:
        logger.error(f"删除智能体 collection 失败: {e}")
        return {"status": "error", "message": f"删除知识库失败: {str(e)}"}


def _get_export_dir(session_id: str = "") -> str:
    """获取导出文件保存目录（按会话隔离，独立于知识库文档目录）

    导出文件保存到 data/export/{session_id}/ 目录，与知识库文档 (data/documents/) 分离。
    这样下载路由可以明确区分「知识库原始文档」和「AI生成导出文档」。
    按会话子目录存放，实现会话级清理，删除某会话时只删该会话的导出文件。

    Args:
        session_id: 会话ID，为空时返回 export 根目录（兼容旧逻辑）
    """
    if session_id:
        export_dir = os.path.join(settings.DATA_DIR, "export", session_id)
    else:
        export_dir = os.path.join(settings.DATA_DIR, "export")
    os.makedirs(export_dir, exist_ok=True)
    return export_dir


def export_document_as_docx(content: str, filename: str, title: str = "", session_id: str = "") -> dict:
    """
    将文本内容导出为 .docx 文件，保存到专用导出目录，供用户下载

    整体格式优化：
    - Markdown表格 → Word原生表格（智能列宽、表头样式、单元格内自动换行）
    - 中文序号标题自动识别（一、二、→ H2，1. 2. → H3 等）
    - Markdown标题 / 列表 / 粗体 / 斜体 → Word 对应样式
    - 过滤 === --- 分隔线，消除多余空行

    Args:
        content: 文档内容（纯文本或Markdown格式）
        filename: 输出文件名（含扩展名）
        title: 文档标题（可选，将作为文档第一行加粗显示）
        session_id: 会话ID（用于按会话隔离导出文件目录，删除会话时只删该会话的文件）

    Returns:
        dict: 包含导出状态和文件路径
    """
    export_dir = _get_export_dir(session_id=session_id)

    # URL解码文件名：LLM传入的文件名可能是URL编码的（如 %E8%AE%BE...），统一解码为中文
    filename = unquote(filename)
    # 安全文件名：去除路径分隔符，防止路径穿越
    safe_filename = filename.replace('/', '_').replace('\\', '_')
    filename = safe_filename

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # 回退方案：保存为 .txt
        txt_filename = filename.rsplit('.', 1)[0] + '.txt'
        file_path = os.path.join(export_dir, txt_filename)
        with open(file_path, "w", encoding="utf-8") as f:
            if title:
                f.write(f"{title}\n{'=' * len(title)}\n\n")
            f.write(content)
        return {
            "status": "success",
            "filename": txt_filename,
            "file_path": file_path,
            "message": f"文档已导出为 {txt_filename}（python-docx 未安装，回退为 txt 格式）",
        }

    try:
        file_path = os.path.join(export_dir, filename)
        doc = DocxDocument()

        # ===== 设置页面边距（更紧凑） =====
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # ===== 设置默认字体和段落间距 =====
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        pf = style.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15

        # ===== 设置标题样式 =====
        for level in range(1, 5):
            heading_style_name = f'Heading {level}'
            if heading_style_name in doc.styles:
                hs = doc.styles[heading_style_name]
                hs.font.name = '宋体'
                hs.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                hs.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
                hs.paragraph_format.space_after = Pt(4)

        # 添加文档标题
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc_title = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ')
            heading = doc.add_heading(doc_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== 中文序号标题识别模式 =====
        # 一、二、三、 → H2    四级以上中文序号 → H3
        _CN_HEADING_H2 = re.compile(r'^[一二三四五六七八九十]+、\s*.+')
        # （一）（二） → H3
        _CN_HEADING_H3 = re.compile(r'^[（(][一二三四五六七八九十]+[）)]\s*.+')
        # 1. 2. 3. （数字+点） → 已由有序列表处理，不重复
        # 第X章/节 → H2/H3
        _CN_HEADING_CHAPTER = re.compile(r'^第[一二三四五六七八九十\d]+[章节部篇]\s*.+')

        # ===== 解析内容并写入 Word =====
        lines = content.split('\n')
        i = 0
        _need_portrait_after_table = False
        while i < len(lines):
            line_stripped = lines[i].strip()

            # 1) 空行 → 跳过
            if not line_stripped:
                i += 1
                continue

            # 2) 分隔线（=== --- *** 等）→ 跳过，不写入Word
            if re.match(r'^[=\-*]{5,}$', line_stripped):
                i += 1
                continue

            # 3) Markdown 表格检测
            if line_stripped.startswith('|') and '|' in line_stripped[1:]:
                table_lines = []
                while i < len(lines):
                    row_line = lines[i].strip()
                    if row_line.startswith('|') and '|' in row_line[1:]:
                        table_lines.append(row_line)
                        i += 1
                    else:
                        break

                # 解析表格数据（跳过分隔行 |---|---|）
                parsed_rows = []
                separator_skipped = False
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    is_separator = all(re.match(r'^:?-+:?$', c.strip()) for c in cells if c.strip())
                    if is_separator and not separator_skipped:
                        separator_skipped = True
                        continue
                    parsed_rows.append(cells)

                if parsed_rows:
                    num_cols = max(len(row) for row in parsed_rows)
                    for row in parsed_rows:
                        while len(row) < num_cols:
                            row.append('')

                    num_rows = len(parsed_rows)

                    # ===== 宽表格（>8列）：自动切换横向页面 =====
                    if num_cols > 8:
                        _add_landscape_section(doc, title_text=f"（续）横向表格 - {num_cols}列")

                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # ===== 智能列宽计算与设置 =====
                    col_widths = _calc_table_col_widths(parsed_rows, num_cols, table)
                    _apply_table_col_widths(table, col_widths)

                    # ===== 宽表格后切换回纵向 =====
                    _need_portrait_after_table = num_cols > 8

                    # 填充表格内容
                    for row_idx, row_data in enumerate(parsed_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            cell_text = _clean_markdown_formatting(cell_text).strip()
                            cell = table.cell(row_idx, col_idx)
                            cell.text = ''
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            p = cell.paragraphs[0]

                            if row_idx == 0:
                                # 表头行：加粗 + 浅蓝背景
                                run = p.add_run(cell_text)
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.name = '宋体'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                shading = OxmlElement('w:shd')
                                shading.set(qn('w:fill'), 'D9E2F3')
                                shading.set(qn('w:val'), 'clear')
                                cell._tc.get_or_add_tcPr().append(shading)
                            else:
                                # 数据行：富文本（支持粗体/斜体）
                                _add_rich_run_to_paragraph(p, cell_text, font_size=Pt(10))

                            # 单元格内段落间距紧凑
                            p.paragraph_format.space_before = Pt(1)
                            p.paragraph_format.space_after = Pt(1)
                            p.paragraph_format.line_spacing = 1.0

                    # 表格后加一个紧凑间隔
                    spacer = doc.add_paragraph('')
                    spacer.paragraph_format.space_before = Pt(2)
                    spacer.paragraph_format.space_after = Pt(2)

                    # ===== 宽表格后切换回纵向页面 =====
                    if _need_portrait_after_table:
                        _add_portrait_section(doc)
                        _need_portrait_after_table = False
                continue

            # 4) Markdown 标题检测（从长到短匹配，避免误判）
            heading_level = None
            heading_text = None
            if line_stripped.startswith('##### '):
                heading_level, heading_text = 5, line_stripped[6:].strip()
            elif line_stripped.startswith('#### '):
                heading_level, heading_text = 4, line_stripped[5:].strip()
            elif line_stripped.startswith('### '):
                heading_level, heading_text = 3, line_stripped[4:].strip()
            elif line_stripped.startswith('## '):
                heading_level, heading_text = 2, line_stripped[3:].strip()
            elif line_stripped.startswith('# ') and not line_stripped.startswith('## '):
                heading_level, heading_text = 2, line_stripped[2:].strip()

            if heading_level is not None:
                heading_text = _clean_markdown_formatting(heading_text)
                # Word heading最高4级
                doc.add_heading(heading_text, level=min(heading_level, 4))
                i += 1
                continue

            # 5) 中文序号标题识别
            if _CN_HEADING_CHAPTER.match(line_stripped):
                text = _clean_markdown_formatting(line_stripped)
                doc.add_heading(text, level=2)
                i += 1
                continue
            if _CN_HEADING_H2.match(line_stripped):
                text = _clean_markdown_formatting(line_stripped)
                doc.add_heading(text, level=2)
                i += 1
                continue
            if _CN_HEADING_H3.match(line_stripped):
                text = _clean_markdown_formatting(line_stripped)
                doc.add_heading(text, level=3)
                i += 1
                continue

            # 6) 列表项
            if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                text = line_stripped[2:].strip()
                text = _clean_markdown_formatting(text)
                _add_rich_paragraph(doc, text, style='List Bullet')
                i += 1
                continue
            if re.match(r'^\d+\.\s', line_stripped):
                text = re.sub(r'^\d+\.\s+', '', line_stripped)
                text = _clean_markdown_formatting(text)
                _add_rich_paragraph(doc, text, style='List Number')
                i += 1
                continue

            # 7) 普通段落（支持行内粗体/斜体）
            text = _clean_markdown_formatting(line_stripped)
            _add_rich_paragraph(doc, text)
            i += 1

        doc.save(file_path)
        logger.info(f"[导出] 文档已生成: {file_path}")

        return {
            "status": "success",
            "filename": filename,
            "file_path": file_path,
            "message": f"文档已导出为 {filename}",
        }
    except Exception as e:
        logger.error(f"[导出] docx生成失败: {e}", exc_info=True)
        txt_filename = filename.rsplit('.', 1)[0] + '.txt'
        file_path = os.path.join(export_dir, txt_filename)
        with open(file_path, "w", encoding="utf-8") as f:
            if title:
                f.write(f"{title}\n\n")
            f.write(content)
        return {
            "status": "success",
            "filename": txt_filename,
            "file_path": file_path,
            "message": f"文档已导出为 {txt_filename}（docx 生成失败: {str(e)}，回退为 txt 格式）",
        }


def export_document_as_xlsx(content: str, filename: str, title: str = "", session_id: str = "") -> dict:
    """将文本内容导出为 .xlsx 文件，保存到专用导出目录，供用户下载
    
    支持 Markdown 表格语法自动转为 Excel 原生表格。
    
    Args:
        content: 文档内容（Markdown格式，支持表格）
        filename: 输出文件名（含扩展名）
        title: 文档标题（可选，将作为第一个工作表的名称）
        session_id: 会话ID（用于按会话隔离导出文件目录）
    
    Returns:
        dict: 包含导出状态和文件路径
    """
    export_dir = _get_export_dir(session_id=session_id)
    
    filename = unquote(filename)
    safe_filename = filename.replace('/', '_').replace('\\', '_')
    filename = safe_filename
    
    if not filename.endswith('.xlsx'):
        filename = filename.rsplit('.', 1)[0] + '.xlsx'
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        # 回退：保存为 CSV
        csv_filename = filename.rsplit('.', 1)[0] + '.csv'
        file_path = os.path.join(export_dir, csv_filename)
        with open(file_path, "w", encoding="utf-8-sig") as f:
            if title:
                f.write(f"{title}\n\n")
            f.write(content)
        return {
            "status": "success",
            "filename": csv_filename,
            "file_path": file_path,
            "message": f"文档已导出为 {csv_filename}（openpyxl 未安装，回退为 CSV 格式）",
        }
    
    try:
        file_path = os.path.join(export_dir, filename)
        wb = Workbook()
        
        # Styles
        header_font = Font(bold=True, size=11, name='宋体')
        header_fill = PatternFill(start_color='D9E2F3', end_color='D9E2F3', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell_font = Font(size=10, name='宋体')
        cell_alignment = Alignment(vertical='center', wrap_text=True)
        info_font = Font(size=10, name='宋体')
        info_alignment = Alignment(vertical='center', wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Parse content and extract tables
        lines = content.split('\n')
        sheet_index = 0
        current_sheet_name = title or filename.rsplit('.', 1)[0]
        current_rows = []  # list of list of strings
        pending_info_lines = []  # non-table content to be written above the next table
        has_table = False
        
        for line in lines:
            stripped = line.strip()
            
            # Detect Markdown table row
            if stripped.startswith('|') and '|' in stripped[1:]:
                cells = [c.strip() for c in stripped.strip('|').split('|')]
                # Skip separator row |---|---|
                if all(re.match(r'^:?-+:?$', c.strip()) for c in cells if c.strip()):
                    continue
                current_rows.append(cells)
                has_table = True
            else:
                # If we were collecting table rows and now hit non-table line,
                # flush the current table to a sheet
                if current_rows:
                    sheet_index = _write_rows_to_xlsx_sheet(
                        wb, current_rows, sheet_index, current_sheet_name,
                        header_font, header_fill, header_alignment, 
                        cell_font, cell_alignment, thin_border,
                        info_lines=pending_info_lines, info_font=info_font, info_alignment=info_alignment
                    )
                    current_rows = []
                    pending_info_lines = []
                    current_sheet_name = f'表格{sheet_index + 1}'
                
                # Collect non-table content
                if stripped:
                    # Check for heading patterns that could be sheet names
                    heading_match = re.match(r'^#{1,3}\s+(.+)', stripped)
                    if heading_match and not has_table:
                        current_sheet_name = heading_match.group(1).strip()[:31]  # Excel sheet name max 31 chars
                    # Clean markdown formatting and add as pending info
                    clean_text = re.sub(r'\*+', '', stripped)  # Remove markdown bold markers
                    clean_text = re.sub(r'^#{1,6}\s+', '', clean_text)  # Remove heading markers
                    if clean_text.strip():
                        pending_info_lines.append(clean_text.strip())
        
        # Flush remaining table rows
        if current_rows:
            sheet_index = _write_rows_to_xlsx_sheet(
                wb, current_rows, sheet_index, current_sheet_name,
                header_font, header_fill, header_alignment,
                cell_font, cell_alignment, thin_border,
                info_lines=pending_info_lines, info_font=info_font, info_alignment=info_alignment
            )
            pending_info_lines = []
        
        # If there's non-table content but no tables, write text to first sheet
        if pending_info_lines and not has_table:
            ws = wb.active
            ws.title = current_sheet_name[:31] if current_sheet_name else 'Sheet1'
            for row_idx, text in enumerate(pending_info_lines, 1):
                ws.cell(row=row_idx, column=1, value=text)
        
        # Remove default empty sheet if we created others
        if sheet_index > 0 and 'Sheet' in wb.sheetnames:
            del wb['Sheet']
        
        wb.save(file_path)
        
        return {
            "status": "success",
            "filename": filename,
            "file_path": file_path,
            "message": f"文档已导出为 {filename}",
        }
    except Exception as e:
        logger.error(f"导出 XLSX 失败: {e}", exc_info=True)
        return {"status": "error", "message": f"导出 XLSX 失败: {str(e)}"}


def _write_rows_to_xlsx_sheet(wb, rows_data, sheet_index, sheet_name, 
                                header_font, header_fill, header_alignment,
                                cell_font, cell_alignment, thin_border,
                                info_lines=None, info_font=None, info_alignment=None):
    """Write parsed table rows to an XLSX worksheet
    
    Args:
        info_lines: Optional list of non-table text lines to write above the table
                    (e.g., project info like "项目名称: XXX")
        info_font: Font for info lines
        info_alignment: Alignment for info lines
    
    Returns:
        int: next sheet index
    """
    if sheet_index == 0:
        ws = wb.active
        ws.title = (sheet_name or '表格1')[:31]
    else:
        ws = wb.create_sheet(title=(sheet_name or f'表格{sheet_index+1}')[:31])
    
    # Write info lines above the table (e.g., project metadata)
    info_row_count = 0
    if info_lines:
        for i, text in enumerate(info_lines):
            row_num = i + 1
            # Try to parse "key: value" or "key：value" patterns into two columns
            kv_match = re.match(r'^(.+?)[：:]\s*(.+)$', text)
            if kv_match:
                cell_key = ws.cell(row=row_num, column=1, value=kv_match.group(1).strip())
                cell_val = ws.cell(row=row_num, column=2, value=kv_match.group(2).strip())
                cell_key.font = info_font or cell_font
                cell_key.alignment = info_alignment or cell_alignment
                cell_val.font = info_font or cell_font
                cell_val.alignment = info_alignment or cell_alignment
            else:
                cell = ws.cell(row=row_num, column=1, value=text)
                cell.font = info_font or cell_font
                cell.alignment = info_alignment or cell_alignment
            info_row_count = i + 1
    
    # Add a blank row between info and table if info exists
    table_start_row = info_row_count + 1 if info_row_count > 0 else 1
    
    num_cols = max(len(row) for row in rows_data)
    
    for row_idx, row in enumerate(rows_data, table_start_row):
        for col_idx, cell_text in enumerate(row[:num_cols], 1):
            # Clean markdown formatting
            clean_text = re.sub(r'\*+', '', cell_text)
            cell = ws.cell(row=row_idx, column=col_idx, value=clean_text)
            cell.border = thin_border
            
            if row_idx == table_start_row:
                # Header row
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            else:
                cell.font = cell_font
                cell.alignment = cell_alignment
    
    # Auto-adjust column widths (CJK-aware: 中文算2单位，英文算1单位)
    # [BUG FIX] 原代码用 len() 一个中文和一个英文都算1，导致列宽严重不足
    def _display_width(s: str) -> int:
        """计算字符串在 Excel 中的显示宽度（中文≈2，英文≈1）"""
        w = 0
        for ch in str(s):
            if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef':
                w += 2
            else:
                w += 1
        return w
    
    # [BUG FIX] 自动合并多行表头：第一行宽分类列数 < 第二行子列数时，合并对应单元格
    if len(rows_data) >= 2 and num_cols > 1:
        row1_content_cols = sum(1 for c in rows_data[0] if str(c).strip())
        if row1_content_cols > 0 and row1_content_cols < num_cols:
            col = 1
            while col <= num_cols:
                if col <= len(rows_data[0]) and str(rows_data[0][col - 1]).strip():
                    start = col
                    col += 1
                    while col <= num_cols:
                        if col > len(rows_data[0]) or not str(rows_data[0][col - 1]).strip():
                            col += 1
                        else:
                            break
                    end = col - 1
                    if end > start:
                        ws.merge_cells(start_row=table_start_row, start_column=start,
                                       end_row=table_start_row, end_column=end)
                    start = col
                col += 1
    
    for col_idx in range(1, num_cols + 1):
        max_length = 0
        for row in rows_data:
            if col_idx - 1 < len(row):
                cell_len = _display_width(str(row[col_idx - 1]))
                if cell_len > max_length:
                    max_length = cell_len
        # Also consider info_lines width for columns 1-2
        if info_lines and col_idx <= 2:
            for text in info_lines:
                kv_match = re.match(r'^(.+?)[：:]\s*(.+)$', text)
                if kv_match and col_idx == 1:
                    max_length = max(max_length, _display_width(kv_match.group(1).strip()))
                elif kv_match and col_idx == 2:
                    max_length = max(max_length, _display_width(kv_match.group(2).strip()))
                elif col_idx == 1:
                    max_length = max(max_length, _display_width(text))
        # [BUG FIX] 宽表场景（>15列）放宽上限到80，避免中文内容挤成一团
        cap = 80 if num_cols > 15 else 50
        adjusted_width = min(max(max_length + 3, 8), cap)
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = adjusted_width
    
    return sheet_index + 1


def _add_landscape_section(doc, title_text: str = ""):
    """在文档中插入一个横向（Landscape）页面分节符

    宽表格（>8列）在纵向A4页面上放不下，自动切换为横向页面。
    横向A4页面内容区宽度约23cm，比纵向16cm多7cm。

    Args:
        doc: python-docx Document 对象
        title_text: 可选的说明文字
    """
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm as _Cm

    # 添加分节符（新节）
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    # 横向A4: 宽29.7cm, 高21cm
    new_section.page_width = _Cm(29.7)
    new_section.page_height = _Cm(21.0)
    new_section.left_margin = _Cm(2.0)
    new_section.right_margin = _Cm(2.0)
    new_section.top_margin = _Cm(2.0)
    new_section.bottom_margin = _Cm(2.0)


def _add_portrait_section(doc):
    """在文档中插入一个纵向（Portrait）页面分节符，恢复默认排版

    宽表格结束后切回纵向页面。

    Args:
        doc: python-docx Document 对象
    """
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm as _Cm

    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.PORTRAIT
    # 纵向A4: 宽21cm, 高29.7cm
    new_section.page_width = _Cm(21.0)
    new_section.page_height = _Cm(29.7)
    new_section.left_margin = _Cm(2.5)
    new_section.right_margin = _Cm(2.5)
    new_section.top_margin = _Cm(2.0)
    new_section.bottom_margin = _Cm(2.0)


def _calc_table_col_widths(parsed_rows: list, num_cols: int, table) -> list:
    """根据各列内容长度智能计算表格列宽（使用dxa/twips单位）

    核心逻辑：
    - 2列"键-值"表：自动分配 28%:72%
    - 3列表格：如果是"序号-项目-内容"模式，8%:22%:70%
    - 窄表格（4-8列）：按内容长度加权分配
    - 宽表格（>8列）：最小列宽 1.5cm，允许超出页面（配合横向页面）
    - 超宽表格（>12列）：建议拆分为子表，列宽按内容加权但保证最小可读宽度

    Args:
        parsed_rows: 解析后的表格行数据
        num_cols: 列数
        table: docx Table 对象（用于设置XML）

    Returns:
        list: 各列宽度（dxa单位，1cm = 567dxa）
    """
    # A4横向页面可用宽度约23cm = 13041 dxa
    # A4纵向页面可用宽度约16cm = 9072 dxa
    # 对于宽表格使用横向宽度
    if num_cols > 8:
        total_dxa = 13041  # A4横向 23cm
    else:
        total_dxa = 9072   # A4纵向 16cm

    if num_cols == 0:
        return []

    # 计算各列内容的最大字符长度
    col_max_lens = [0] * num_cols
    for row in parsed_rows:
        for ci in range(min(len(row), num_cols)):
            col_max_lens[ci] = max(col_max_lens[ci], len(row[ci]))

    # 2列"键-值"表特殊处理：短标签 vs 长内容
    if num_cols == 2:
        first_col_len = col_max_lens[0]
        second_col_len = col_max_lens[1]
        # 如果第一列明显短于第二列（标签-值模式），分配 28%:72%
        if first_col_len < second_col_len * 0.7:
            return [int(total_dxa * 0.28), int(total_dxa * 0.72)]
        # 否则按内容长度加权
        total_chars = max(first_col_len + second_col_len, 1)
        ratio1 = first_col_len / total_chars
        ratio1 = max(0.2, min(0.5, ratio1))  # 限制在20%-50%之间
        return [int(total_dxa * ratio1), int(total_dxa * (1 - ratio1))]

    # 3列表格：如果是"序号-项目-内容"模式
    if num_cols == 3:
        first_len = col_max_lens[0]
        second_len = col_max_lens[1]
        third_len = col_max_lens[2]
        # 如果第一列很短（序号列）
        if first_len <= 4 and second_len < third_len * 0.6:
            return [int(total_dxa * 0.08), int(total_dxa * 0.22), int(total_dxa * 0.70)]
        # 按内容长度加权
        total_chars = max(first_len + second_len + third_len, 1)
        r1 = max(0.1, min(0.3, first_len / total_chars))
        r2 = max(0.15, min(0.4, second_len / total_chars))
        r3 = 1.0 - r1 - r2
        return [int(total_dxa * r1), int(total_dxa * r2), int(total_dxa * r3)]

    # 宽表格（>8列）：按内容加权，但保证最小可读宽度
    # 最小列宽：1.5cm = 851 dxa（约5个汉字）
    min_col_dxa = 851 if num_cols <= 12 else 567  # >12列时最小1cm

    total_chars = max(sum(col_max_lens), 1)
    raw_ratios = [max(l, 5) / total_chars for l in col_max_lens]
    # 归一化
    raw_sum = sum(raw_ratios)
    ratios = [r / raw_sum for r in raw_ratios]

    # 计算初始列宽
    col_widths = [int(total_dxa * r) for r in ratios]

    # 确保每列至少达到最小可读宽度
    for ci in range(num_cols):
        if col_widths[ci] < min_col_dxa:
            col_widths[ci] = min_col_dxa

    # 如果总宽度超出页面，允许超宽（横向页面通常能容纳）
    # 不强制缩回到页面宽度，避免列太窄
    total_width = sum(col_widths)
    if total_width > total_dxa * 1.5:
        # 如果超宽太多（1.5倍以上），按比例缩放到1.5倍页面宽度
        scale = (total_dxa * 1.5) / total_width
        col_widths = [int(w * scale) for w in col_widths]

    return col_widths


def _apply_table_col_widths(table, col_widths_dxa: list):
    """将列宽应用到Word表格（通过XML确保生效）

    Args:
        table: docx Table 对象
        col_widths_dxa: 各列宽度（dxa单位）
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    total_dxa = sum(col_widths_dxa)
    num_cols = len(col_widths_dxa)

    # 1. 设置 tblGrid（列宽定义）
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gc in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(gc)
    else:
        tblGrid = OxmlElement('w:tblGrid')
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblPr.append(tblGrid)

    for w in col_widths_dxa:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)

    # 2. 设置表格总宽度
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        for old_tw in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old_tw)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_dxa))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

    # 3. 设置每个单元格的 tcW
    for row in table.rows:
        for ci in range(min(num_cols, len(row.cells))):
            w = col_widths_dxa[ci]
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            for old_w in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old_w)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def _add_rich_paragraph(doc, text: str, style: str = None):
    """添加支持行内粗体/斜体的段落

    Args:
        doc: Document对象
        text: 段落文本（可能含 **粗体** 和 *斜体* 标记）
        style: 段落样式名（如 'List Bullet', 'List Number'），None为普通段落
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()

    _add_rich_run_to_paragraph(p, text, font_size=Pt(11))


def _add_rich_run_to_paragraph(p, text: str, font_size=None):
    """向段落中添加富文本Run（支持 **粗体** 和 *斜体*）

    Args:
        p: Paragraph对象
        text: 文本（可能含 **粗体** 和 *斜体* 标记）
        font_size: 字号，None则使用默认
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if font_size is None:
        font_size = Pt(11)

    # 按 **...** 分割文本，交替设置粗体/普通
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = font_size
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif part:
            # 处理行内斜体 *...*
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and len(ip) > 2:
                    run = p.add_run(ip[1:-1])
                    run.italic = True
                    run.font.size = font_size
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                elif ip:
                    run = p.add_run(ip)
                    run.font.size = font_size
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _clean_markdown_formatting(text: str) -> str:
    """清理Markdown格式标记，转为纯文本（保留粗体/斜体标记给 _add_rich_paragraph 使用）"""
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def cleanup_export_files(session_id: str = "", username: str = "") -> int:
    """清理导出文件目录

    按会话粒度清理：只删除属于指定 session_id 的导出文件。
    导出文件存放在 data/export/{session_id}/ 子目录中，实现会话隔离。

    当 session_id 为空时，清理超过 24 小时的过期导出文件（兜底策略）。

    Args:
        session_id: 会话ID（为空时清理过期文件）
        username: 用户名（预留，暂未使用）

    Returns:
        int: 删除的文件数量
    """
    import shutil  # [BUG FIX] 原代码用 shutil.rmtree 但函数内未导入
    export_dir = _get_export_dir()
    if not os.path.exists(export_dir):
        return 0

    deleted_count = 0

    if session_id:
        # 按会话清理：删除 data/export/{session_id}/ 整个目录
        session_export_dir = os.path.join(export_dir, session_id)
        if os.path.exists(session_export_dir) and os.path.isdir(session_export_dir):
            try:
                file_count = len([f for f in os.listdir(session_export_dir) if os.path.isfile(os.path.join(session_export_dir, f))])
                shutil.rmtree(session_export_dir)
                deleted_count = file_count
                logger.info(f"[导出清理] 已清理会话 {session_id} 的 {file_count} 个导出文件")
            except Exception as e:
                logger.warning(f"[导出清理] 清理会话 {session_id} 导出目录失败: {e}")
    else:
        # 无指定会话：清理超过 24 小时的过期导出文件（兜底，防止文件堆积）
        now = time.time()
        max_age = 86400  # 24 小时
        for item in os.listdir(export_dir):
            item_path = os.path.join(export_dir, item)
            if os.path.isdir(item_path):
                # 会话子目录：检查目录修改时间
                try:
                    dir_mtime = os.path.getmtime(item_path)
                    if now - dir_mtime > max_age:
                        file_count = len([f for f in os.listdir(item_path) if os.path.isfile(os.path.join(item_path, f))])
                        shutil.rmtree(item_path)
                        deleted_count += file_count
                        logger.info(f"[导出清理] 已清理过期会话目录: {item} ({file_count} 个文件)")
                except Exception as e:
                    logger.warning(f"[导出清理] 清理过期目录失败 {item}: {e}")
            elif os.path.isfile(item_path):
                # 兼容旧版：直接放在 export/ 下的文件（无子目录），按文件修改时间清理
                try:
                    file_mtime = os.path.getmtime(item_path)
                    if now - file_mtime > max_age:
                        os.remove(item_path)
                        deleted_count += 1
                        logger.info(f"[导出清理] 已清理过期导出文件: {item}")
                except Exception as e:
                    logger.warning(f"[导出清理] 删除过期文件失败 {item}: {e}")

    return deleted_count


def list_all_collections() -> list[dict]:
    """列出 ChromaDB 中所有的 collection 及其文档数（诊断用）"""
    import chromadb
    try:
        client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
        collections = client.list_collections()
        result = []
        for c in collections:
            try:
                count = c.count()
            except:
                count = -1
            result.append({"name": c.name, "count": count})

        # [#11] 同时列出关键词索引信息
        keyword_info = []
        if os.path.exists(KEYWORD_INDEX_DIR):
            for fname in os.listdir(KEYWORD_INDEX_DIR):
                if fname.startswith("index_") and fname.endswith(".json"):
                    fpath = os.path.join(KEYWORD_INDEX_DIR, fname)
                    try:
                        with open(fpath, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        keyword_info.append({"name": fname, "count": len(data), "type": "keyword"})
                    except Exception:
                        pass

        result.extend(keyword_info)
        return result
    except Exception as e:
        return [{"name": "error", "count": 0, "message": str(e)}]
